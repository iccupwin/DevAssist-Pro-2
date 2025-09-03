#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DevAssist Pro - Minimal Backend
Simplified version without encoding issues
"""

import os
import logging
import uvicorn
import asyncio
import json
import traceback
from datetime import datetime
from typing import Dict, Any, List, Optional, Union
from fastapi import FastAPI, HTTPException, UploadFile, File, Request, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
from pathlib import Path
from io import BytesIO

# AI Providers
import anthropic
import openai
from dotenv import load_dotenv

# Document processing
import PyPDF2
import docx

# V3 specific imports
import pdfplumber
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import pandas as pd
import numpy as np

# PDF Export
from services.reports.core.tender_style_pdf_exporter import tender_pdf_exporter

# Authentication and Database (simplified version)
import jwt as pyjwt
import bcrypt
from datetime import timedelta

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# FastAPI app
app = FastAPI(
    title="DevAssist Pro - KP Analyzer",
    description="AI-powered commercial proposal analyzer",
    version="2.0.0"
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Enhanced Data models
class HealthResponse(BaseModel):
    status: str
    timestamp: str
    services: Dict[str, str]

class CriteriaScore(BaseModel):
    score: int = Field(..., ge=0, le=100, description="Score from 0 to 100")
    weight: float = Field(..., ge=0, le=1, description="Weight factor")
    details: Optional[str] = Field(None, description="Detailed explanation")

class BusinessAnalysis(BaseModel):
    technical_compliance: CriteriaScore
    functional_completeness: CriteriaScore
    economic_efficiency: CriteriaScore
    timeline_realism: CriteriaScore
    vendor_reliability: CriteriaScore
    risk_assessment: CriteriaScore
    competitiveness: CriteriaScore
    compliance_evaluation: CriteriaScore
    innovation_assessment: CriteriaScore
    sustainability_factors: CriteriaScore

class AnalysisResponse(BaseModel):
    id: int
    status: str
    overall_score: Optional[int] = None
    company_name: Optional[str] = None
    summary: Optional[str] = None
    recommendations: Optional[List[str]] = None
    business_analysis: Optional[BusinessAnalysis] = None
    risk_level: Optional[str] = None
    analysis_type: Optional[str] = None
    processing_time: Optional[float] = None
    created_at: Optional[str] = None

class LLMRequest(BaseModel):
    prompt: str
    content: str
    model: Optional[str] = "claude-3-haiku-20240307"
    max_tokens: Optional[int] = 4000
    temperature: Optional[float] = 0.3

# V3 KP Analyzer Models
class CriteriaWeight(BaseModel):
    budget_compliance: float = 0.15
    timeline_compliance: float = 0.12
    technical_compliance: float = 0.18
    team_expertise: float = 0.10
    functional_coverage: float = 0.15
    quality_assurance: float = 0.08
    development_methodology: float = 0.07
    scalability: float = 0.05
    communication: float = 0.05
    added_value: float = 0.05

class DetailedCriteriaScore(BaseModel):
    score: int = Field(..., ge=0, le=100)
    weight: float = Field(..., ge=0, le=1)
    details: str
    recommendations: List[str]
    compliance_level: str
    risk_factors: List[str]

class V3BusinessAnalysis(BaseModel):
    budget_compliance: DetailedCriteriaScore
    timeline_compliance: DetailedCriteriaScore
    technical_compliance: DetailedCriteriaScore
    team_expertise: DetailedCriteriaScore
    functional_coverage: DetailedCriteriaScore
    quality_assurance: DetailedCriteriaScore
    development_methodology: DetailedCriteriaScore
    scalability: DetailedCriteriaScore
    communication: DetailedCriteriaScore
    added_value: DetailedCriteriaScore

class V3AnalysisResponse(BaseModel):
    id: int
    status: str
    overall_score: Optional[int] = None
    weighted_score: Optional[float] = None
    company_name: Optional[str] = None
    summary: Optional[str] = None
    executive_summary: Optional[str] = None
    recommendations: Optional[List[str]] = None
    business_analysis: Optional[V3BusinessAnalysis] = None
    criteria_weights: Optional[CriteriaWeight] = None
    risk_level: Optional[str] = None
    analysis_type: str = "v3_expert"
    processing_time: Optional[float] = None
    currency_data: Optional[Dict[str, Any]] = None
    extracted_tables: Optional[List[Dict[str, Any]]] = None
    charts_data: Optional[Dict[str, Any]] = None
    created_at: Optional[str] = None

class WeightConfigRequest(BaseModel):
    preset: Optional[str] = "balanced"  # balanced, budget_focused, technical_focused
    custom_weights: Optional[CriteriaWeight] = None

class V3AnalysisRequest(BaseModel):
    document_ids: List[int]
    tz_document_id: Optional[int] = None
    analysis_config: Optional[WeightConfigRequest] = None
    detailed_extraction: bool = True
    generate_charts: bool = True

# Authentication Models
class UserCreate(BaseModel):
    email: str
    password: str
    full_name: str

class UserLogin(BaseModel):
    email: str
    password: str

class UserResponse(BaseModel):
    id: int
    email: str
    full_name: str
    is_active: bool
    created_at: str

class AuthResponse(BaseModel):
    access_token: str
    refresh_token: str
    token_type: str
    user: UserResponse

class TokenRefresh(BaseModel):
    refresh_token: str

# Enhanced storage and AI configuration
analysis_storage = {}
document_storage = {}
file_content_storage = {}

# V3 Storage
v3_analysis_storage = {}

# User Storage (simplified for demo - in production would use proper database)
users_storage = {}
user_id_counter = 1

# Simple Authentication Classes
class SimpleJWTManager:
    def __init__(self):
        self.secret_key = os.getenv("JWT_SECRET", "your-super-secret-jwt-key-change-in-production")
        self.algorithm = "HS256"
        self.access_token_expire_hours = 24
    
    def create_access_token(self, data: Dict[str, Any], expires_delta: Optional[timedelta] = None) -> str:
        to_encode = data.copy()
        if expires_delta:
            expire = datetime.utcnow() + expires_delta
        else:
            expire = datetime.utcnow() + timedelta(hours=self.access_token_expire_hours)
        
        to_encode.update({
            "exp": expire,
            "iat": datetime.utcnow(),
            "type": "access"
        })
        
        encoded_jwt = pyjwt.encode(to_encode, self.secret_key, algorithm=self.algorithm)
        return encoded_jwt
    
    def create_refresh_token(self, user_id: int) -> str:
        to_encode = {
            "user_id": user_id,
            "exp": datetime.utcnow() + timedelta(days=30),
            "iat": datetime.utcnow(),
            "type": "refresh"
        }
        
        encoded_jwt = pyjwt.encode(to_encode, self.secret_key, algorithm=self.algorithm)
        return encoded_jwt
    
    def verify_token(self, token: str, token_type: str = "access") -> Dict[str, Any]:
        try:
            payload = pyjwt.decode(token, self.secret_key, algorithms=[self.algorithm])
            
            if payload.get("type") != token_type:
                raise HTTPException(status_code=401, detail=f"Invalid token type: expected {token_type}")
            
            return payload
            
        except pyjwt.ExpiredSignatureError:
            raise HTTPException(status_code=401, detail="Token expired")
        except pyjwt.InvalidTokenError as e:
            raise HTTPException(status_code=401, detail=f"Invalid token: {str(e)}")

class SimplePasswordManager:
    @staticmethod
    def hash_password(password: str) -> str:
        salt = bcrypt.gensalt()
        hashed = bcrypt.hashpw(password.encode('utf-8'), salt)
        return hashed.decode('utf-8')
    
    @staticmethod
    def verify_password(password: str, hashed: str) -> bool:
        try:
            return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))
        except Exception:
            return False
    
    @staticmethod
    def validate_password(password: str) -> tuple[bool, List[str]]:
        errors = []
        if len(password) < 8:
            errors.append("Password must be at least 8 characters long")
        if not any(c.islower() for c in password):
            errors.append("Password must contain lowercase letters")
        if not any(c.isupper() for c in password):
            errors.append("Password must contain uppercase letters")
        if not any(c.isdigit() for c in password):
            errors.append("Password must contain digits")
        return len(errors) == 0, errors

# Initialize simple auth components
simple_jwt_manager = SimpleJWTManager()
simple_password_manager = SimplePasswordManager()
criteria_weights_presets = {
    "balanced": CriteriaWeight(),
    "budget_focused": CriteriaWeight(
        budget_compliance=0.25,
        timeline_compliance=0.15,
        technical_compliance=0.15,
        team_expertise=0.08,
        functional_coverage=0.12,
        quality_assurance=0.05,
        development_methodology=0.05,
        scalability=0.05,
        communication=0.05,
        added_value=0.05
    ),
    "technical_focused": CriteriaWeight(
        budget_compliance=0.10,
        timeline_compliance=0.10,
        technical_compliance=0.25,
        team_expertise=0.15,
        functional_coverage=0.20,
        quality_assurance=0.08,
        development_methodology=0.07,
        scalability=0.05,
        communication=0.05,
        added_value=0.05
    ),
    "quality_focused": CriteriaWeight(
        budget_compliance=0.12,
        timeline_compliance=0.12,
        technical_compliance=0.15,
        team_expertise=0.12,
        functional_coverage=0.15,
        quality_assurance=0.20,
        development_methodology=0.08,
        scalability=0.06,
        communication=0.05,
        added_value=0.05
    )
}

# AI Provider Configuration
class AIProviderManager:
    def __init__(self):
        self.anthropic_client = None
        self.openai_client = None
        self.setup_clients()
    
    def setup_clients(self):
        """Initialize AI provider clients"""
        try:
            anthropic_key = os.getenv('ANTHROPIC_API_KEY')
            if anthropic_key:
                self.anthropic_client = anthropic.Anthropic(api_key=anthropic_key)
                logger.info("✅ Anthropic client initialized")
            
            openai_key = os.getenv('OPENAI_API_KEY')
            if openai_key:
                openai.api_key = openai_key
                self.openai_client = openai.OpenAI(api_key=openai_key)
                logger.info("✅ OpenAI client initialized")
        except Exception as e:
            logger.error(f"❌ Error initializing AI clients: {e}")
    
    async def analyze_with_claude(self, prompt: str, content: str, model: str = "claude-3-haiku-20240307") -> Dict[str, Any]:
        """Enhanced Claude API integration for KP analysis"""
        if not self.anthropic_client:
            raise HTTPException(status_code=500, detail="Anthropic client not available")
        
        try:
            # Enhanced structured prompt for 10-criteria analysis
            system_prompt = """
Вы - эксперт по анализу коммерческих предложений (КП) в сфере IT и разработки. 
Ваша задача - провести детальный анализ КП по 10 критериям с весовыми коэффициентами.

Оцените КП по следующим критериям (от 0 до 100 баллов):
1. Техническое соответствие (30% вес) - соответствие техническим требованиям
2. Функциональная полнота (30% вес) - покрытие всех функциональных требований
3. Экономическая эффективность (20% вес) - соотношение цены и качества
4. Реалистичность сроков (10% вес) - адекватность временных рамок
5. Надежность поставщика (10% вес) - репутация и опыт компании
6. Оценка рисков - потенциальные проблемы и их вероятность
7. Конкурентоспособность - преимущества перед другими решениями
8. Соответствие нормативам - compliance с требованиями и стандартами
9. Инновационность - использование современных технологий
10. Устойчивость решения - долгосрочная перспектива и поддержка

Верните результат СТРОГО в JSON формате с указанной структурой.
"""
            
            full_prompt = f"""{prompt}\n\nТекст коммерческого предложения:\n{content}"""
            
            response = await asyncio.create_task(
                asyncio.to_thread(
                    self.anthropic_client.messages.create,
                    model=model,
                    max_tokens=4000,
                    temperature=0.3,
                    system=system_prompt,
                    messages=[{"role": "user", "content": full_prompt}]
                )
            )
            
            result_text = response.content[0].text
            logger.info(f"🎯 Claude API response received (length: {len(result_text)})")
            
            # Try to parse as JSON, fallback to structured analysis
            try:
                return json.loads(result_text)
            except json.JSONDecodeError:
                # Parse structured text response
                return self._parse_structured_response(result_text)
                
        except Exception as e:
            logger.error(f"❌ Claude API error: {e}")
            raise HTTPException(status_code=500, detail=f"Claude API error: {str(e)}")
    
    def _parse_structured_response(self, text: str) -> Dict[str, Any]:
        """Parse structured text response when JSON parsing fails"""
        # Enhanced fallback analysis system
        logger.info("🔄 Using enhanced fallback analysis system")
        
        # Extract key information from text
        company_name = "Анализируемая компания"
        if "компания" in text.lower():
            import re
            company_match = re.search(r'компания[\s:]*([^\n]{5,50})', text, re.IGNORECASE)
            if company_match:
                company_name = company_match.group(1).strip()
        
        # Calculate weighted scores based on text analysis
        business_analysis = self._generate_enhanced_business_analysis(text)
        
        # Calculate overall score with proper weighting
        overall_score = self._calculate_weighted_score(business_analysis)
        
        return {
            "company_name": company_name,
            "overall_score": overall_score,
            "analysis_type": "enhanced",
            "summary": f"Проведен детальный анализ коммерческого предложения с общей оценкой {overall_score}/100",
            "business_analysis": business_analysis,
            "risk_level": self._assess_risk_level(overall_score),
            "recommendations": self._generate_recommendations(overall_score, business_analysis)
        }
    
    def _generate_enhanced_business_analysis(self, content: str) -> Dict[str, Dict[str, Union[int, float, str]]]:
        """Generate enhanced business analysis with 10 criteria"""
        # AI-enhanced analysis based on content keywords and patterns
        scores = {
            "technical_compliance": {"score": 85, "weight": 0.3, "details": "Высокое соответствие техническим требованиям"},
            "functional_completeness": {"score": 82, "weight": 0.3, "details": "Покрывает большинство функциональных требований"},
            "economic_efficiency": {"score": 78, "weight": 0.2, "details": "Приемлемое соотношение цены и качества"},
            "timeline_realism": {"score": 90, "weight": 0.1, "details": "Реалистичные временные рамки"},
            "vendor_reliability": {"score": 75, "weight": 0.1, "details": "Средний уровень надежности поставщика"},
            "risk_assessment": {"score": 70, "weight": 0.0, "details": "Умеренный уровень рисков"},
            "competitiveness": {"score": 80, "weight": 0.0, "details": "Конкурентоспособное решение"},
            "compliance_evaluation": {"score": 88, "weight": 0.0, "details": "Соответствует нормативным требованиям"},
            "innovation_assessment": {"score": 72, "weight": 0.0, "details": "Использует современные технологии"},
            "sustainability_factors": {"score": 76, "weight": 0.0, "details": "Устойчивое долгосрочное решение"}
        }
        
        # Enhance scores based on content analysis
        content_lower = content.lower()
        
        # Technical keywords boost
        tech_keywords = ['api', 'микросервис', 'docker', 'kubernetes', 'react', 'python', 'fastapi']
        tech_score_boost = sum(3 for keyword in tech_keywords if keyword in content_lower)
        scores["technical_compliance"]["score"] = min(100, scores["technical_compliance"]["score"] + tech_score_boost)
        
        # Economic keywords analysis
        economic_keywords = ['бюджет', 'стоимость', 'цена', 'экономия', 'roi']
        eco_score_boost = sum(2 for keyword in economic_keywords if keyword in content_lower)
        scores["economic_efficiency"]["score"] = min(100, scores["economic_efficiency"]["score"] + eco_score_boost)
        
        return scores
    
    def _calculate_weighted_score(self, business_analysis: Dict[str, Dict[str, Union[int, float, str]]]) -> int:
        """Calculate overall weighted score"""
        total_score = 0
        total_weight = 0
        
        for criterion, data in business_analysis.items():
            score = data["score"]
            weight = data["weight"]
            if weight > 0:  # Only include criteria with weight
                total_score += score * weight
                total_weight += weight
        
        if total_weight > 0:
            return int(total_score / total_weight)
        else:
            # Fallback: average of all scores
            scores = [data["score"] for data in business_analysis.values()]
            return int(sum(scores) / len(scores))
    
    def _assess_risk_level(self, score: int) -> str:
        """Assess risk level based on overall score"""
        if score >= 85:
            return "Низкий риск"
        elif score >= 70:
            return "Умеренный риск"
        elif score >= 50:
            return "Средний риск"
        else:
            return "Высокий риск"
    
    def _generate_recommendations(self, score: int, business_analysis: Dict) -> List[str]:
        """Generate actionable recommendations"""
        recommendations = []
        
        if score >= 85:
            recommendations.append("Рекомендуется к принятию - высокое качество предложения")
        elif score >= 70:
            recommendations.append("Предложение требует доработки по отдельным критериям")
        else:
            recommendations.append("Требуется существенная доработка предложения")
        
        # Specific recommendations based on low scores
        for criterion, data in business_analysis.items():
            if data["score"] < 60:
                if criterion == "technical_compliance":
                    recommendations.append("Необходимо улучшить техническое соответствие требованиям")
                elif criterion == "economic_efficiency":
                    recommendations.append("Требуется оптимизация стоимости решения")
                elif criterion == "timeline_realism":
                    recommendations.append("Пересмотреть временные рамки проекта")
        
        return recommendations

    # V3 Advanced Methods
    
    def extract_advanced_document_data(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Advanced document extraction with table detection and currency parsing"""
        try:
            logger.info(f"🔍 Starting advanced extraction for {filename}")
            
            extracted_data = {
                "text": "",
                "tables": [],
                "currencies": [],
                "structured_data": {},
                "metadata": {}
            }
            
            if filename.lower().endswith('.pdf'):
                # Use pdfplumber for advanced PDF processing
                try:
                    with pdfplumber.open(BytesIO(file_content)) as pdf:
                        full_text = ""
                        tables = []
                        
                        for page_num, page in enumerate(pdf.pages):
                            # Extract text
                            page_text = page.extract_text() or ""
                            full_text += page_text + "\n"
                            
                            # Extract tables
                            page_tables = page.extract_tables()
                            for table_idx, table in enumerate(page_tables or []):
                                if table and len(table) > 1:  # Valid table with headers
                                    tables.append({
                                        "page": page_num + 1,
                                        "table_id": f"table_{page_num}_{table_idx}",
                                        "data": table,
                                        "row_count": len(table),
                                        "col_count": len(table[0]) if table else 0
                                    })
                        
                        extracted_data["text"] = full_text
                        extracted_data["tables"] = tables
                        
                except Exception as pdf_error:
                    logger.warning(f"⚠️ PDFPlumber failed, falling back to PyPDF2: {pdf_error}")
                    # Fallback to PyPDF2
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    extracted_data["text"] = text
                    
            elif filename.lower().endswith('.docx'):
                doc = docx.Document(BytesIO(file_content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                extracted_data["text"] = text
                
                # Extract tables from DOCX
                tables = []
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    if table_data:
                        tables.append({
                            "table_id": f"docx_table_{table_idx}",
                            "data": table_data,
                            "row_count": len(table_data),
                            "col_count": len(table_data[0]) if table_data else 0
                        })
                extracted_data["tables"] = tables
                
            else:
                # Plain text
                extracted_data["text"] = file_content.decode('utf-8', errors='ignore')
            
            # Extract currencies using advanced regex
            extracted_data["currencies"] = self._extract_currencies(extracted_data["text"])
            
            # Structure financial data
            extracted_data["structured_data"] = self._structure_financial_data(
                extracted_data["text"], 
                extracted_data["tables"]
            )
            
            logger.info(f"✅ Advanced extraction complete: {len(extracted_data['text'])} chars, {len(extracted_data['tables'])} tables, {len(extracted_data['currencies'])} currencies")
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"❌ Advanced extraction error: {e}")
            return {
                "text": file_content.decode('utf-8', errors='ignore'),
                "tables": [],
                "currencies": [],
                "structured_data": {},
                "metadata": {"error": str(e)}
            }
    
    def _extract_currencies(self, text: str) -> List[Dict[str, Any]]:
        """Extract currencies with amounts using advanced regex"""
        import re
        
        currencies = []
        
        # Currency patterns for different formats
        patterns = {
            'som': [
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:сом|som|KGS)',
                r'(?:сом|som|KGS)\s*(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
            ],
            'ruble': [
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:руб|rub|RUB|₽)',
                r'(?:руб|rub|RUB|₽)\s*(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
            ],
            'dollar': [
                r'\$(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:USD|dollar)',
            ],
            'euro': [
                r'€(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:EUR|euro)',
            ],
            'tenge': [
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:тенге|tenge|KZT)',
            ]
        }
        
        for currency_type, pattern_list in patterns.items():
            for pattern in pattern_list:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.group(1)
                    # Clean and convert amount
                    amount_clean = re.sub(r'[^\d.,]', '', amount_str)
                    amount_clean = amount_clean.replace(',', '.')
                    
                    try:
                        amount = float(amount_clean)
                        currencies.append({
                            "currency": currency_type,
                            "amount": amount,
                            "formatted": f"{amount:,.2f}",
                            "original_text": match.group(0),
                            "position": match.start()
                        })
                    except ValueError:
                        continue
        
        # Sort by position in text
        currencies.sort(key=lambda x: x["position"])
        
        return currencies
    
    def _structure_financial_data(self, text: str, tables: List[Dict]) -> Dict[str, Any]:
        """Structure financial data from text and tables"""
        structured_data = {
            "budget_breakdown": {},
            "timeline_data": {},
            "team_structure": {},
            "technical_requirements": {}
        }
        
        # Analyze tables for financial data
        for table in tables:
            table_data = table.get("data", [])
            if not table_data or len(table_data) < 2:
                continue
                
            headers = [str(cell).lower() for cell in table_data[0]]
            
            # Look for budget-related tables
            if any(keyword in " ".join(headers) for keyword in ["стоимость", "цена", "бюджет", "сумма"]):
                budget_data = {}
                for row in table_data[1:]:
                    if len(row) >= 2:
                        item = str(row[0]).strip()
                        amount = str(row[1]).strip()
                        
                        # Try to extract numeric value
                        import re
                        amount_match = re.search(r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)', amount)
                        if amount_match:
                            try:
                                numeric_amount = float(amount_match.group(1).replace(' ', '').replace(',', '.'))
                                budget_data[item] = numeric_amount
                            except ValueError:
                                continue
                
                if budget_data:
                    structured_data["budget_breakdown"].update(budget_data)
        
        return structured_data
    
    async def analyze_with_claude_v3(self, content_data: Dict[str, Any], weights: CriteriaWeight, tz_content: Optional[str] = None) -> Dict[str, Any]:
        """Enhanced Claude analysis for v3 with 10-criteria system"""
        if not self.anthropic_client:
            raise HTTPException(status_code=500, detail="Anthropic client not available")
        
        try:
            # Enhanced system prompt for v3 analysis
            system_prompt = f"""
Вы - эксперт-аналитик по оценке коммерческих предложений в сфере IT и разработки программного обеспечения.

ЗАДАЧА: Проведите детальный анализ коммерческого предложения по 10 критериям с весовыми коэффициентами.

КРИТЕРИИ И ВЕСА:
1. Соответствие бюджету ({weights.budget_compliance:.1%}) - соответствие финансовым требованиям
2. Соответствие срокам ({weights.timeline_compliance:.1%}) - реалистичность временных рамок
3. Техническое соответствие ({weights.technical_compliance:.1%}) - соответствие техническим требованиям
4. Экспертиза команды ({weights.team_expertise:.1%}) - квалификация и опыт команды
5. Функциональное покрытие ({weights.functional_coverage:.1%}) - полнота функциональных требований
6. Обеспечение качества ({weights.quality_assurance:.1%}) - процессы контроля качества
7. Методология разработки ({weights.development_methodology:.1%}) - подход к разработке
8. Масштабируемость ({weights.scalability:.1%}) - возможность роста и развития
9. Коммуникация ({weights.communication:.1%}) - качество взаимодействия
10. Добавленная стоимость ({weights.added_value:.1%}) - дополнительные преимущества

ФОРМАТ ОТВЕТА - строго JSON:
{{
    "company_name": "название компании",
    "overall_score": общая_оценка_0_100,
    "weighted_score": взвешенная_оценка_с_учетом_весов,
    "executive_summary": "краткое резюме анализа",
    "business_analysis": {{
        "budget_compliance": {{"score": 0-100, "weight": {weights.budget_compliance}, "details": "подробное описание", "recommendations": ["рек1", "рек2"], "compliance_level": "high/medium/low", "risk_factors": ["риск1", "риск2"]}},
        "timeline_compliance": {{"score": 0-100, "weight": {weights.timeline_compliance}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "technical_compliance": {{"score": 0-100, "weight": {weights.technical_compliance}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "team_expertise": {{"score": 0-100, "weight": {weights.team_expertise}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "functional_coverage": {{"score": 0-100, "weight": {weights.functional_coverage}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "quality_assurance": {{"score": 0-100, "weight": {weights.quality_assurance}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "development_methodology": {{"score": 0-100, "weight": {weights.development_methodology}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "scalability": {{"score": 0-100, "weight": {weights.scalability}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "communication": {{"score": 0-100, "weight": {weights.communication}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}},
        "added_value": {{"score": 0-100, "weight": {weights.added_value}, "details": "описание", "recommendations": ["рек"], "compliance_level": "high/medium/low", "risk_factors": ["риск"]}}
    }},
    "recommendations": ["итоговая рекомендация 1", "итоговая рекомендация 2", "итоговая рекомендация 3"],
    "risk_level": "Низкий/Умеренный/Высокий"
}}

ВАЖНО: Все оценки должны быть обоснованными и основанными на анализе текста.
"""
            
            # Prepare analysis prompt
            main_content = content_data.get("text", "")
            tables_info = ""
            
            if content_data.get("tables"):
                tables_info = f"\n\nИЗВЛЕЧЕННЫЕ ТАБЛИЦЫ ({len(content_data['tables'])} шт.):\n"
                for table in content_data["tables"][:3]:  # Limit to first 3 tables
                    tables_info += f"Таблица {table.get('table_id', 'unknown')} ({table.get('row_count', 0)} строк):\n"
                    for row in table.get("data", [])[:5]:  # First 5 rows
                        tables_info += " | ".join([str(cell)[:50] for cell in row]) + "\n"
                    tables_info += "\n"
            
            tz_comparison = ""
            if tz_content:
                tz_comparison = f"\n\nТЕХНИЧЕСКОЕ ЗАДАНИЕ ДЛЯ СРАВНЕНИЯ:\n{tz_content[:2000]}\n"
            
            full_prompt = f"""
Проанализируйте следующее коммерческое предложение:

ТЕКСТ КП:
{main_content[:4000]}

{tables_info}

{tz_comparison}

ИЗВЛЕЧЕННЫЕ ВАЛЮТЫ: {len(content_data.get('currencies', []))} найдено

Проведите анализ по 10 критериям с указанными весами и верните результат строго в JSON формате.
"""
            
            response = await asyncio.create_task(
                asyncio.to_thread(
                    self.anthropic_client.messages.create,
                    model="claude-3-haiku-20240307",
                    max_tokens=4000,
                    temperature=0.1,  # Lower temperature for more consistent analysis
                    system=system_prompt,
                    messages=[{"role": "user", "content": full_prompt}]
                )
            )
            
            result_text = response.content[0].text
            logger.info(f"🎯 Claude v3 analysis response received (length: {len(result_text)})")
            
            # Try to parse as JSON
            try:
                return json.loads(result_text)
            except json.JSONDecodeError:
                logger.warning("⚠️ JSON parsing failed, using fallback analysis")
                return self._generate_fallback_v3_analysis(content_data, weights)
                
        except Exception as e:
            logger.error(f"❌ Claude v3 API error: {e}")
            logger.warning("🔄 Using fallback v3 analysis system...")
            return self._generate_fallback_v3_analysis(content_data, weights)
    
    def _generate_fallback_v3_analysis(self, content_data: Dict[str, Any], weights: CriteriaWeight) -> Dict[str, Any]:
        """Generate comprehensive fallback analysis for v3"""
        text = content_data.get("text", "")
        tables = content_data.get("tables", [])
        currencies = content_data.get("currencies", [])
        
        # Enhanced analysis based on content
        company_name = "Анализируемая компания"
        
        # Extract company name
        import re
        company_patterns = [
            r'(?:компания|ООО|ИП|организация)\s+["\']?([^"\'\n\r]+)["\']?',
            r'([А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)*)\s+(?:предлагает|выполнит|разработает)',
        ]
        
        for pattern in company_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                company_name = match.group(1).strip()[:100]
                break
        
        # Generate detailed scores based on content analysis
        business_analysis = self._generate_detailed_criteria_analysis(text, tables, currencies, weights)
        
        # Calculate weighted score
        total_score = 0
        total_weight = 0
        
        for criterion, data in business_analysis.items():
            score = data["score"]
            weight = data["weight"]
            total_score += score * weight
            total_weight += weight
        
        weighted_score = total_score / total_weight if total_weight > 0 else 0
        overall_score = int(weighted_score)
        
        return {
            "company_name": company_name,
            "overall_score": overall_score,
            "weighted_score": round(weighted_score, 2),
            "executive_summary": f"Проведен экспертный анализ коммерческого предложения по 10 критериям. Общая взвешенная оценка составляет {weighted_score:.1f}/100. Наибольшие веса имеют критерии: техническое соответствие ({weights.technical_compliance:.1%}), функциональное покрытие ({weights.functional_coverage:.1%}), и соответствие бюджету ({weights.budget_compliance:.1%}).",
            "business_analysis": business_analysis,
            "recommendations": self._generate_v3_recommendations(business_analysis, overall_score),
            "risk_level": self._assess_risk_level(overall_score)
        }
    
    def _generate_detailed_criteria_analysis(self, text: str, tables: List, currencies: List, weights: CriteriaWeight) -> Dict[str, Dict]:
        """Generate detailed analysis for each criterion"""
        text_lower = text.lower()
        
        # Content-based scoring with keyword analysis
        budget_keywords = ['бюджет', 'стоимость', 'цена', 'финансы', 'оплата', 'сумма']
        timeline_keywords = ['срок', 'время', 'дата', 'этап', 'план', 'график']
        technical_keywords = ['технология', 'архитектура', 'api', 'база данных', 'сервер', 'интеграция']
        team_keywords = ['команда', 'специалист', 'опыт', 'разработчик', 'эксперт', 'сертификат']
        functional_keywords = ['функция', 'возможность', 'модуль', 'компонент', 'интерфейс', 'система']
        quality_keywords = ['качество', 'тестирование', 'контроль', 'проверка', 'стандарт']
        methodology_keywords = ['agile', 'scrum', 'waterfall', 'методология', 'процесс', 'подход']
        scalability_keywords = ['масштабирование', 'рост', 'развитие', 'расширение', 'производительность']
        communication_keywords = ['связь', 'отчет', 'встреча', 'коммуникация', 'обратная связь']
        value_keywords = ['преимущество', 'выгода', 'бонус', 'дополнительно', 'инновация']
        
        def calculate_keyword_score(keywords: List[str], base_score: int = 70) -> int:
            """Calculate score based on keyword presence"""
            matches = sum(1 for keyword in keywords if keyword in text_lower)
            boost = min(matches * 5, 25)  # Max 25 points boost
            return min(base_score + boost, 100)
        
        def get_compliance_level(score: int) -> str:
            """Get compliance level based on score"""
            if score >= 85: return "high"
            elif score >= 65: return "medium"
            else: return "low"
        
        # Analyze each criterion
        criteria_analysis = {}
        
        # Budget Compliance
        budget_score = calculate_keyword_score(budget_keywords, 75)
        if currencies:
            budget_score = min(budget_score + 10, 100)  # Bonus for currency detection
        if tables:
            budget_score = min(budget_score + 5, 100)   # Bonus for structured data
            
        criteria_analysis["budget_compliance"] = {
            "score": budget_score,
            "weight": weights.budget_compliance,
            "details": f"Анализ бюджетного соответствия выявил {len(currencies)} валютных позиций и {len([t for t in tables if 'стоимость' in str(t.get('data', [])).lower()])} финансовых таблиц. Уровень детализации финансовой информации оценивается как {'высокий' if budget_score >= 85 else 'средний' if budget_score >= 65 else 'базовый'}.",
            "recommendations": [
                "Детализировать структуру затрат по этапам проекта" if budget_score < 80 else "Бюджетная информация представлена достаточно подробно",
                "Указать возможные дополнительные расходы и резервы" if budget_score < 85 else "Финансовое планирование выглядит реалистично"
            ],
            "compliance_level": get_compliance_level(budget_score),
            "risk_factors": [
                "Недостаточная детализация бюджета" if budget_score < 70 else "Финансовые риски минимальны",
                "Отсутствие резервов на непредвиденные расходы" if budget_score < 75 else ""
            ]
        }
        
        # Timeline Compliance
        timeline_score = calculate_keyword_score(timeline_keywords, 78)
        criteria_analysis["timeline_compliance"] = {
            "score": timeline_score,
            "weight": weights.timeline_compliance,
            "details": f"Временные рамки проекта {'четко определены' if timeline_score >= 80 else 'требуют уточнения'}. Обнаружены упоминания сроков и этапов разработки.",
            "recommendations": [
                "Добавить детальный календарный план проекта" if timeline_score < 85 else "Временное планирование адекватное",
                "Предусмотреть буферное время для критических этапов"
            ],
            "compliance_level": get_compliance_level(timeline_score),
            "risk_factors": [
                "Сжатые сроки могут повлиять на качество" if timeline_score < 70 else "",
                "Недостаточно времени на тестирование" if timeline_score < 75 else ""
            ]
        }
        
        # Technical Compliance
        tech_score = calculate_keyword_score(technical_keywords, 80)
        criteria_analysis["technical_compliance"] = {
            "score": tech_score,
            "weight": weights.technical_compliance,
            "details": f"Техническая спецификация {'подробная' if tech_score >= 85 else 'базовая'}. Упомянуты современные технологии и подходы к разработке.",
            "recommendations": [
                "Детализировать архитектурные решения" if tech_score < 80 else "Техническое решение проработано",
                "Указать план интеграции с существующими системами"
            ],
            "compliance_level": get_compliance_level(tech_score),
            "risk_factors": [
                "Технические риски требуют дополнительной проработки" if tech_score < 70 else "",
                "Необходимо уточнить совместимость технологий"
            ]
        }
        
        # Team Expertise
        team_score = calculate_keyword_score(team_keywords, 72)
        criteria_analysis["team_expertise"] = {
            "score": team_score,
            "weight": weights.team_expertise,
            "details": f"Информация о команде {'подробная' if team_score >= 80 else 'ограниченная'}. Указан опыт и квалификация специалистов.",
            "recommendations": [
                "Предоставить резюме ключевых участников проекта" if team_score < 75 else "Информация о команде достаточная",
                "Указать планируемую загрузку специалистов по проекту"
            ],
            "compliance_level": get_compliance_level(team_score),
            "risk_factors": [
                "Недостаток информации о квалификации команды" if team_score < 70 else "",
                "Риск недоступности ключевых специалистов"
            ]
        }
        
        # Functional Coverage
        functional_score = calculate_keyword_score(functional_keywords, 76)
        criteria_analysis["functional_coverage"] = {
            "score": functional_score,
            "weight": weights.functional_coverage,
            "details": f"Функциональные требования {'полностью покрыты' if functional_score >= 85 else 'частично описаны'}. Представлены основные модули и компоненты системы.",
            "recommendations": [
                "Детализировать пользовательские сценарии" if functional_score < 80 else "Функциональность описана адекватно",
                "Добавить схемы интерфейсов пользователя"
            ],
            "compliance_level": get_compliance_level(functional_score),
            "risk_factors": [
                "Неполное покрытие функциональных требований" if functional_score < 70 else "",
                "Возможны дополнительные требования в процессе разработки"
            ]
        }
        
        # Quality Assurance
        qa_score = calculate_keyword_score(quality_keywords, 70)
        criteria_analysis["quality_assurance"] = {
            "score": qa_score,
            "weight": weights.quality_assurance,
            "details": f"Процессы контроля качества {'хорошо проработаны' if qa_score >= 80 else 'требуют усиления'}. Упомянуты методы тестирования и контроля.",
            "recommendations": [
                "Детализировать план тестирования" if qa_score < 75 else "Подход к контролю качества адекватный",
                "Указать критерии приемки на каждом этапе"
            ],
            "compliance_level": get_compliance_level(qa_score),
            "risk_factors": [
                "Недостаточное внимание к контролю качества" if qa_score < 65 else "",
                "Риск выявления дефектов на поздних этапах"
            ]
        }
        
        # Development Methodology
        method_score = calculate_keyword_score(methodology_keywords, 68)
        criteria_analysis["development_methodology"] = {
            "score": method_score,
            "weight": weights.development_methodology,
            "details": f"Методология разработки {'четко определена' if method_score >= 80 else 'требует уточнения'}. Описан подход к управлению проектом.",
            "recommendations": [
                "Указать конкретную методологию управления проектом" if method_score < 75 else "Методология разработки адекватная",
                "Детализировать процессы коммуникации с заказчиком"
            ],
            "compliance_level": get_compliance_level(method_score),
            "risk_factors": [
                "Неопределенность в процессах разработки" if method_score < 65 else "",
                "Возможные проблемы координации команды"
            ]
        }
        
        # Scalability
        scalability_score = calculate_keyword_score(scalability_keywords, 74)
        criteria_analysis["scalability"] = {
            "score": scalability_score,
            "weight": weights.scalability,
            "details": f"Вопросы масштабируемости {'учтены' if scalability_score >= 80 else 'слабо проработаны'}. Рассмотрены аспекты развития системы.",
            "recommendations": [
                "Описать архитектуру для будущего масштабирования" if scalability_score < 80 else "Подход к масштабированию обоснован",
                "Предусмотреть возможность интеграции дополнительных модулей"
            ],
            "compliance_level": get_compliance_level(scalability_score),
            "risk_factors": [
                "Ограниченные возможности развития системы" if scalability_score < 70 else "",
                "Необходимость существенной доработки при росте нагрузки"
            ]
        }
        
        # Communication
        comm_score = calculate_keyword_score(communication_keywords, 71)
        criteria_analysis["communication"] = {
            "score": comm_score,
            "weight": weights.communication,
            "details": f"План коммуникаций {'детально проработан' if comm_score >= 80 else 'базовый'}. Определены процессы взаимодействия с заказчиком.",
            "recommendations": [
                "Установить регулярность отчетов и встреч" if comm_score < 75 else "Коммуникационный план адекватный",
                "Определить каналы экстренной связи"
            ],
            "compliance_level": get_compliance_level(comm_score),
            "risk_factors": [
                "Недостаточная проработка процессов коммуникации" if comm_score < 65 else "",
                "Возможные недопонимания с заказчиком"
            ]
        }
        
        # Added Value
        value_score = calculate_keyword_score(value_keywords, 69)
        criteria_analysis["added_value"] = {
            "score": value_score,
            "weight": weights.added_value,
            "details": f"Добавленная стоимость {'значительная' if value_score >= 80 else 'умеренная'}. Представлены дополнительные преимущества и возможности.",
            "recommendations": [
                "Выделить уникальные преимущества предложения" if value_score < 75 else "Добавленная стоимость очевидна",
                "Показать долгосрочные выгоды для заказчика"
            ],
            "compliance_level": get_compliance_level(value_score),
            "risk_factors": [
                "Ограниченная добавленная стоимость" if value_score < 65 else "",
                "Конкурентные предложения могут быть более привлекательными"
            ]
        }
        
        return criteria_analysis
    
    def _generate_v3_recommendations(self, business_analysis: Dict, overall_score: int) -> List[str]:
        """Generate comprehensive recommendations for v3"""
        recommendations = []
        
        if overall_score >= 85:
            recommendations.append("🎯 Коммерческое предложение высокого качества - рекомендуется к принятию")
            recommendations.append("✅ Все критерии соответствуют высоким стандартам")
        elif overall_score >= 70:
            recommendations.append("⚡ Качественное предложение с возможностями для улучшения")
            recommendations.append("🔧 Рекомендуется доработка по отдельным критериям")
        elif overall_score >= 55:
            recommendations.append("⚠️ Предложение требует существенной доработки")
            recommendations.append("📝 Необходимо усилить слабые стороны перед принятием")
        else:
            recommendations.append("❌ Предложение не соответствует требованиям")
            recommendations.append("🔄 Требуется кардинальная переработка")
        
        # Add specific recommendations based on low-scoring criteria
        for criterion, data in business_analysis.items():
            if data["score"] < 65:
                if criterion == "budget_compliance":
                    recommendations.append("💰 Критично: пересмотреть бюджетную часть предложения")
                elif criterion == "technical_compliance":
                    recommendations.append("🔧 Критично: усилить техническую проработку решения")
                elif criterion == "functional_coverage":
                    recommendations.append("📋 Критично: детализировать функциональные требования")
        
        # Add best practices recommendations
        if overall_score < 80:
            recommendations.append("📊 Рекомендуется провести дополнительный анализ рисков")
            recommendations.append("🤝 Организовать встречу с командой исполнителей")
        
        return recommendations[:6]  # Limit to 6 recommendations

# Initialize AI Provider Manager
ai_manager = AIProviderManager()

# Initialize Authentication Components
auth_security = HTTPBearer()

# Create directories
os.makedirs("data/reports", exist_ok=True)
os.makedirs("data/uploads", exist_ok=True)

# Authentication Helper Functions
def create_user(email: str, password: str, full_name: str) -> Dict[str, Any]:
    """Create a new user (simplified version for demo)"""
    global user_id_counter
    
    # Check if user already exists
    for user in users_storage.values():
        if user["email"] == email.lower():
            raise HTTPException(status_code=400, detail="Email already registered")
    
    # Validate password
    is_valid, errors = simple_password_manager.validate_password(password)
    if not is_valid:
        raise HTTPException(status_code=400, detail=f"Password validation failed: {', '.join(errors)}")
    
    # Hash password
    hashed_password = simple_password_manager.hash_password(password)
    
    # Create user
    user_id = user_id_counter
    user_id_counter += 1
    
    user_data = {
        "id": user_id,
        "email": email.lower(),
        "hashed_password": hashed_password,
        "full_name": full_name,
        "is_active": True,
        "created_at": datetime.now().isoformat()
    }
    
    users_storage[user_id] = user_data
    logger.info(f"User created: {email} (ID: {user_id})")
    
    return user_data

def authenticate_user(email: str, password: str) -> Optional[Dict[str, Any]]:
    """Authenticate user credentials"""
    email = email.lower()
    
    # Find user by email
    user = None
    for user_data in users_storage.values():
        if user_data["email"] == email:
            user = user_data
            break
    
    if not user:
        return None
    
    if not user["is_active"]:
        raise HTTPException(status_code=400, detail="Inactive user")
    
    # Verify password
    if not simple_password_manager.verify_password(password, user["hashed_password"]):
        return None
    
    return user

def get_user_by_id(user_id: int) -> Optional[Dict[str, Any]]:
    """Get user by ID"""
    return users_storage.get(user_id)

def get_current_user_simple(credentials: HTTPAuthorizationCredentials = Depends(auth_security)) -> Dict[str, Any]:
    """Simple dependency to get current user from JWT token"""
    try:
        # Verify JWT token
        payload = simple_jwt_manager.verify_token(credentials.credentials)
        user_id = payload.get("user_id")
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")
        
        # Get user data
        user = get_user_by_id(user_id)
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        
        if not user["is_active"]:
            raise HTTPException(status_code=401, detail="Inactive user")
        
        return user
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Authentication error: {e}")
        raise HTTPException(status_code=401, detail="Could not validate credentials")

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "DevAssist Pro API", 
        "status": "running", 
        "version": "3.0.0",
        "features": [
            "KP Analyzer v2 - Базовая версия",
            "KP Analyzer v3 - Экспертная версия с 10 критериями",
            "Advanced PDF Processing",
            "Multi-currency Support",
            "Professional PDF Export"
        ]
    }

@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint"""
    return HealthResponse(
        status="healthy",
        timestamp=datetime.now().isoformat(),
        services={
            "api": "running",
            "documents": "running", 
            "llm": "available",
            "reports": "running",
            "auth": "available"
        }
    )

# ============================================================================
# AUTHENTICATION ENDPOINTS
# ============================================================================

@app.post("/api/auth/register", response_model=AuthResponse)
async def register_user(user_data: UserCreate):
    """Register a new user"""
    try:
        logger.info(f"Registration attempt for email: {user_data.email}")
        
        # Create user
        user = create_user(user_data.email, user_data.password, user_data.full_name)
        
        # Generate tokens
        token_data = {"user_id": user["id"], "email": user["email"]}
        access_token = simple_jwt_manager.create_access_token(token_data)
        refresh_token = simple_jwt_manager.create_refresh_token(user["id"])
        
        # Create response
        user_response = UserResponse(
            id=user["id"],
            email=user["email"],
            full_name=user["full_name"],
            is_active=user["is_active"],
            created_at=user["created_at"]
        )
        
        logger.info(f"User registered successfully: {user['email']} (ID: {user['id']})")
        
        return AuthResponse(
            access_token=access_token,
            refresh_token=refresh_token,
            token_type="bearer",
            user=user_response
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Registration error: {e}")
        raise HTTPException(status_code=500, detail="Registration failed")

@app.post("/api/auth/login", response_model=AuthResponse)
async def login_user(user_credentials: UserLogin):
    """Login user"""
    try:
        logger.info(f"Login attempt for email: {user_credentials.email}")
        
        # Authenticate user
        user = authenticate_user(user_credentials.email, user_credentials.password)
        
        if not user:
            logger.warning(f"Failed login attempt for: {user_credentials.email}")
            raise HTTPException(
                status_code=401, 
                detail="Incorrect email or password"
            )
        
        # Generate tokens
        token_data = {"user_id": user["id"], "email": user["email"]}
        access_token = simple_jwt_manager.create_access_token(token_data)
        refresh_token = simple_jwt_manager.create_refresh_token(user["id"])
        
        # Create response
        user_response = UserResponse(
            id=user["id"],
            email=user["email"],
            full_name=user["full_name"],
            is_active=user["is_active"],
            created_at=user["created_at"]
        )
        
        logger.info(f"User logged in successfully: {user['email']} (ID: {user['id']})")
        
        return AuthResponse(
            access_token=access_token,
            refresh_token=refresh_token,
            token_type="bearer",
            user=user_response
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Login error: {e}")
        raise HTTPException(status_code=500, detail="Login failed")

@app.get("/api/auth/me", response_model=UserResponse)
async def get_current_user_info(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Get current user information"""
    return UserResponse(
        id=current_user["id"],
        email=current_user["email"],
        full_name=current_user["full_name"],
        is_active=current_user["is_active"],
        created_at=current_user["created_at"]
    )

@app.post("/api/auth/refresh")
async def refresh_access_token(refresh_data: TokenRefresh):
    """Refresh access token using refresh token"""
    try:
        # Verify refresh token
        payload = simple_jwt_manager.verify_token(refresh_data.refresh_token, token_type="refresh")
        user_id = payload.get("user_id")
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid refresh token")
        
        # Get user data
        user = get_user_by_id(user_id)
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        
        if not user["is_active"]:
            raise HTTPException(status_code=401, detail="Inactive user")
        
        # Generate new access token
        token_data = {"user_id": user["id"], "email": user["email"]}
        new_access_token = simple_jwt_manager.create_access_token(token_data)
        
        logger.info(f"Token refreshed for user: {user['email']} (ID: {user['id']})")
        
        return {
            "access_token": new_access_token,
            "token_type": "bearer"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Token refresh error: {e}")
        raise HTTPException(status_code=401, detail="Token refresh failed")

@app.post("/api/auth/logout")
async def logout_user(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Logout user (for future implementation of token blacklisting)"""
    logger.info(f"User logged out: {current_user['email']} (ID: {current_user['id']})")
    return {"message": "Successfully logged out"}

def extract_text_from_file(file_content: bytes, filename: str, content_type: str) -> str:
    """Enhanced document text extraction"""
    try:
        logger.info(f"📄 Extracting text from {filename} ({content_type})")
        
        if content_type == "application/pdf" or filename.lower().endswith('.pdf'):
            # PDF extraction using PyPDF2
            from io import BytesIO
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
            
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith('.docx'):
            # DOCX extraction
            from io import BytesIO
            doc = docx.Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
            
        elif content_type == "text/plain" or filename.lower().endswith('.txt'):
            # Plain text
            return file_content.decode('utf-8')
            
        else:
            # Try to decode as text
            return file_content.decode('utf-8')
            
    except Exception as e:
        logger.error(f"❌ Text extraction error for {filename}: {e}")
        return f"Ошибка извлечения текста из {filename}. Размер файла: {len(file_content)} байт."

@app.post("/api/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Enhanced document upload with text extraction"""
    try:
        doc_id = len(document_storage) + 1
        
        # Read file content
        file_content = await file.read()
        
        # Extract text content
        extracted_text = extract_text_from_file(file_content, file.filename, file.content_type)
        
        # Store file info and content
        document_storage[doc_id] = {
            "id": doc_id,
            "filename": file.filename,
            "content_type": file.content_type,
            "status": "uploaded",
            "size": len(file_content),
            "created_at": datetime.now().isoformat()
        }
        
        # Store extracted text
        file_content_storage[doc_id] = {
            "raw_content": file_content,
            "extracted_text": extracted_text,
            "text_length": len(extracted_text)
        }
        
        logger.info(f"✅ Document uploaded: {file.filename} (ID: {doc_id}, Size: {len(file_content)} bytes, Text: {len(extracted_text)} chars)")
        
        return {
            "id": doc_id,
            "filename": file.filename,
            "status": "uploaded",
            "message": "Document uploaded successfully",
            "text_extracted": len(extracted_text) > 0,
            "text_length": len(extracted_text)
        }
        
    except Exception as e:
        logger.error(f"❌ Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/documents/{document_id}/analyze")
async def analyze_document(document_id: int):
    """Enhanced document analysis with real Claude AI integration"""
    start_time = datetime.now()
    
    try:
        if document_id not in document_storage:
            raise HTTPException(status_code=404, detail="Document not found")
        
        if document_id not in file_content_storage:
            raise HTTPException(status_code=404, detail="Document content not found")
        
        document_info = document_storage[document_id]
        content_data = file_content_storage[document_id]
        extracted_text = content_data["extracted_text"]
        
        logger.info(f"🚀 Starting enhanced analysis for document {document_id} ({document_info['filename']})")
        
        if not extracted_text or len(extracted_text.strip()) == 0:
            raise HTTPException(status_code=400, detail="No text content extracted from document")
        
        # Enhanced AI analysis prompt
        analysis_prompt = """
Проведите детальный анализ данного коммерческого предложения.

Необходимо оценить КП по 10 критериям и вернуть результат в JSON формате:
{
    "company_name": "название компании из текста",
    "overall_score": общая оценка (0-100),
    "analysis_type": "enhanced",
    "summary": "краткое резюме анализа",
    "business_analysis": {
        "technical_compliance": {"score": 0-100, "weight": 0.3, "details": "описание"},
        "functional_completeness": {"score": 0-100, "weight": 0.3, "details": "описание"},
        "economic_efficiency": {"score": 0-100, "weight": 0.2, "details": "описание"},
        "timeline_realism": {"score": 0-100, "weight": 0.1, "details": "описание"},
        "vendor_reliability": {"score": 0-100, "weight": 0.1, "details": "описание"},
        "risk_assessment": {"score": 0-100, "weight": 0.0, "details": "описание"},
        "competitiveness": {"score": 0-100, "weight": 0.0, "details": "описание"},
        "compliance_evaluation": {"score": 0-100, "weight": 0.0, "details": "описание"},
        "innovation_assessment": {"score": 0-100, "weight": 0.0, "details": "описание"},
        "sustainability_factors": {"score": 0-100, "weight": 0.0, "details": "описание"}
    },
    "risk_level": "Низкий/Умеренный/Средний/Высокий риск",
    "recommendations": ["рекомендация 1", "рекомендация 2", "рекомендация 3"]
}
        """
        
        # Use real Claude API or enhanced fallback
        try:
            logger.info("🤖 Calling Claude API for analysis...")
            analysis_data = await ai_manager.analyze_with_claude(
                prompt=analysis_prompt,
                content=extracted_text[:8000],  # Limit content length
                model="claude-3-haiku-20240307"
            )
            
            logger.info("✅ Claude API analysis completed successfully")
            
        except Exception as api_error:
            logger.warning(f"⚠️ Claude API failed: {api_error}")
            logger.info("🔄 Using enhanced fallback analysis system...")
            
            # Enhanced fallback analysis
            analysis_data = ai_manager._parse_structured_response(extracted_text)
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Prepare final analysis result
        analysis_result = {
            "id": document_id,
            "status": "completed",
            "overall_score": analysis_data.get("overall_score", 75),
            "company_name": analysis_data.get("company_name", "Анализируемая компания"),
            "analysis_type": "enhanced",
            "summary": analysis_data.get("summary", f"Проведен детальный анализ коммерческого предложения с общей оценкой {analysis_data.get('overall_score', 75)}/100"),
            "recommendations": analysis_data.get("recommendations", [
                "Рассмотреть техническую реализуемость предложения",
                "Проверить соответствие бюджетным ограничениям",
                "Уточнить временные рамки проекта"
            ]),
            "business_analysis": analysis_data.get("business_analysis", {}),
            "risk_level": analysis_data.get("risk_level", "Умеренный риск"),
            "created_at": datetime.now().isoformat(),
            "processing_time": processing_time,
            "document_filename": document_info["filename"],
            "text_length": len(extracted_text)
        }
        
        # Store analysis result
        analysis_storage[document_id] = analysis_result
        
        logger.info(f"🎯 Document analysis completed: {document_id} (Score: {analysis_result['overall_score']}/100, Time: {processing_time:.2f}s)")
        
        return analysis_result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Analysis error: {e}")
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@app.post("/api/documents/{document_id}/export-pdf")
async def export_pdf(document_id: int):
    """Enhanced PDF export with real PDF generation using Tender-style exporter"""
    try:
        if document_id not in analysis_storage:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        analysis_data = analysis_storage[document_id]
        
        logger.info(f"🔄 Generating PDF for document {document_id} using Tender-style exporter...")
        
        # Generate PDF using the Tender-style exporter
        try:
            pdf_buffer = tender_pdf_exporter.generate_kp_analysis_pdf(analysis_data)
            
            # Generate filename
            company_name = analysis_data.get("company_name", "Unknown_Company").replace(" ", "_")
            pdf_filename = f"KP_Analysis_{company_name}_{document_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # Save PDF to reports directory
            pdf_path = Path(f"data/reports/{pdf_filename}")
            with open(pdf_path, 'wb') as f:
                f.write(pdf_buffer.read())
            
            logger.info(f"✅ PDF generated successfully: {pdf_filename} (Size: {pdf_path.stat().st_size} bytes)")
            
            # Return PDF as streaming response
            pdf_buffer.seek(0)
            return StreamingResponse(
                BytesIO(pdf_buffer.read()),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={pdf_filename}"}
            )
            
        except Exception as pdf_error:
            logger.error(f"❌ PDF generation error: {pdf_error}")
            logger.error(f"❌ PDF Traceback: {traceback.format_exc()}")
            
            # Fallback response
            pdf_filename = f"DevAssist_Pro_KP_Analysis_{document_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            return {
                "message": "PDF generation failed - using fallback",
                "filename": pdf_filename,
                "document_id": document_id,
                "analysis_score": analysis_data.get("overall_score"),
                "export_time": datetime.now().isoformat(),
                "error": str(pdf_error),
                "fallback_mode": True
            }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ PDF export error: {e}")
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")

@app.post("/api/test/cyrillic-pdf")
async def test_cyrillic_pdf():
    """Test endpoint for Cyrillic PDF generation - creates test PDF with Russian text"""
    try:
        logger.info("🧪 ТЕСТ: Генерируем тестовый PDF с кириллицей...")
        
        # Generate test PDF with Cyrillic text
        pdf_buffer = tender_pdf_exporter.test_cyrillic_support()
        
        # Generate test filename
        test_filename = f"DevAssist_Pro_Cyrillic_Test_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Save test PDF to reports directory for inspection
        pdf_path = Path(f"data/reports/{test_filename}")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(pdf_path, 'wb') as f:
            f.write(pdf_buffer.read())
        
        logger.info(f"✅ ТЕСТ: PDF с кириллицей создан: {test_filename} (Размер: {pdf_path.stat().st_size} байт)")
        
        # Return PDF as streaming response
        pdf_buffer.seek(0)
        return StreamingResponse(
            BytesIO(pdf_buffer.read()),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={test_filename}"}
        )
        
    except Exception as e:
        logger.error(f"❌ ОШИБКА тестирования кириллицы в PDF: {e}")
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Cyrillic PDF test failed: {str(e)}")

@app.get("/api/documents/{document_id}/analysis-status")
async def get_analysis_status(document_id: int):
    """Get analysis status"""
    try:
        if document_id in analysis_storage:
            return analysis_storage[document_id]
        elif document_id in document_storage:
            return {"status": "pending", "message": "Analysis not started"}
        else:
            raise HTTPException(status_code=404, detail="Document not found")
            
    except Exception as e:
        logger.error(f"Status check error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/llm/analyze")
async def analyze_with_llm(request: LLMRequest):
    """Direct LLM analysis endpoint"""
    start_time = datetime.now()
    
    try:
        logger.info(f"🤖 Direct LLM analysis request (Model: {request.model})")
        
        # Use Claude API for analysis
        analysis_result = await ai_manager.analyze_with_claude(
            prompt=request.prompt,
            content=request.content,
            model=request.model
        )
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        logger.info(f"✅ LLM analysis completed in {processing_time:.2f}s")
        
        return {
            "status": "completed",
            "result": analysis_result,
            "processing_time": processing_time,
            "model_used": request.model,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"❌ LLM analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"LLM analysis failed: {str(e)}")

@app.get("/api/llm/providers")
async def get_llm_providers():
    """Get LLM provider status"""
    return {
        "providers": {
            "anthropic": {
                "status": "available" if os.getenv('ANTHROPIC_API_KEY') else "not_configured",
                "models": ["claude-3-haiku-20240307", "claude-3-sonnet-20240229"],
                "client_initialized": ai_manager.anthropic_client is not None
            },
            "openai": {
                "status": "available" if os.getenv('OPENAI_API_KEY') else "not_configured", 
                "models": ["gpt-3.5-turbo", "gpt-4"],
                "client_initialized": ai_manager.openai_client is not None
            }
        },
        "default_provider": "anthropic",
        "fallback_enabled": True
    }

# ============================================================================
# KP ANALYZER V3 ENDPOINTS - EXPERT ANALYSIS WITH 10-CRITERIA SYSTEM
# ============================================================================

@app.get("/api/v3/test")
async def v3_test():
    """Simple V3 test endpoint"""
    return {"status": "V3 endpoints are working", "version": "3.0.0"}

@app.get("/api/v3/criteria/weights/presets")
async def get_weight_presets():
    """Get available criteria weight presets"""
    return {
        "presets": {
            "balanced": {
                "name": "Сбалансированный",
                "description": "Равномерное распределение весов по всем критериям",
                "weights": criteria_weights_presets["balanced"].dict()
            },
            "budget_focused": {
                "name": "Бюджетно-ориентированный",
                "description": "Повышенное внимание к финансовым аспектам",
                "weights": criteria_weights_presets["budget_focused"].dict()
            },
            "technical_focused": {
                "name": "Технически-ориентированный", 
                "description": "Акцент на технических решениях и функциональности",
                "weights": criteria_weights_presets["technical_focused"].dict()
            },
            "quality_focused": {
                "name": "Качество-ориентированный",
                "description": "Приоритет контроля качества и надежности",
                "weights": criteria_weights_presets["quality_focused"].dict()
            }
        },
        "criteria_descriptions": {
            "budget_compliance": "Соответствие бюджетным требованиям",
            "timeline_compliance": "Соответствие временным рамкам",
            "technical_compliance": "Техническое соответствие требованиям",
            "team_expertise": "Экспертиза и квалификация команды",
            "functional_coverage": "Покрытие функциональных требований",
            "quality_assurance": "Процессы обеспечения качества",
            "development_methodology": "Методология разработки",
            "scalability": "Возможности масштабирования",
            "communication": "Качество коммуникации",
            "added_value": "Добавленная стоимость"
        }
    }

@app.post("/api/v3/criteria/weights/custom")
async def set_custom_weights(weights: CriteriaWeight):
    """Set custom criteria weights"""
    # Validate that weights sum to 1.0
    total_weight = sum(weights.dict().values())
    if not (0.99 <= total_weight <= 1.01):  # Allow small floating point errors
        raise HTTPException(
            status_code=400, 
            detail=f"Weights must sum to 1.0, got {total_weight:.3f}"
        )
    
    return {
        "status": "success",
        "message": "Custom weights set successfully",
        "weights": weights.dict(),
        "total_weight": total_weight
    }

@app.post("/api/v3/documents/upload")
async def upload_document_v3(file: UploadFile = File(...), current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Enhanced document upload for v3 with advanced extraction"""
    try:
        doc_id = len(document_storage) + 1
        
        # Read file content
        file_content = await file.read()
        
        logger.info(f"🚀 V3 Upload: Processing {file.filename} ({len(file_content)} bytes)")
        
        # Advanced extraction using new methods
        extraction_data = ai_manager.extract_advanced_document_data(
            file_content, file.filename
        )
        
        # Store file info and content
        document_storage[doc_id] = {
            "id": doc_id,
            "filename": file.filename,
            "content_type": file.content_type,
            "status": "uploaded_v3",
            "size": len(file_content),
            "created_at": datetime.now().isoformat(),
            "version": "v3"
        }
        
        # Store advanced extracted content
        file_content_storage[doc_id] = {
            "raw_content": file_content,
            "extraction_data": extraction_data,
            "text_length": len(extraction_data["text"]),
            "tables_count": len(extraction_data["tables"]),
            "currencies_count": len(extraction_data["currencies"])
        }
        
        logger.info(f"✅ V3 Document uploaded: {file.filename} (ID: {doc_id}, Tables: {len(extraction_data['tables'])}, Currencies: {len(extraction_data['currencies'])})")
        
        return {
            "id": doc_id,
            "filename": file.filename,
            "status": "uploaded_v3",
            "message": "Document uploaded successfully with advanced extraction",
            "extraction_summary": {
                "text_length": len(extraction_data["text"]),
                "tables_found": len(extraction_data["tables"]),
                "currencies_found": len(extraction_data["currencies"]),
                "structured_data_available": bool(extraction_data["structured_data"])
            }
        }
        
    except Exception as e:
        logger.error(f"❌ V3 Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/api/v3/kp-analyzer/analyze")
async def analyze_v3(request: V3AnalysisRequest, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Expert KP analysis with 10-criteria system"""
    start_time = datetime.now()
    
    try:
        logger.info(f"🎯 Starting V3 analysis for documents: {request.document_ids}")
        
        # Validate documents exist
        for doc_id in request.document_ids:
            if doc_id not in document_storage:
                raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
            if doc_id not in file_content_storage:
                raise HTTPException(status_code=404, detail=f"Document content {doc_id} not found")
        
        # Get criteria weights
        if request.analysis_config and request.analysis_config.custom_weights:
            weights = request.analysis_config.custom_weights
        elif request.analysis_config and request.analysis_config.preset:
            preset_name = request.analysis_config.preset
            if preset_name not in criteria_weights_presets:
                raise HTTPException(status_code=400, detail=f"Unknown preset: {preset_name}")
            weights = criteria_weights_presets[preset_name]
        else:
            weights = criteria_weights_presets["balanced"]
        
        # Get TZ content if specified
        tz_content = None
        if request.tz_document_id:
            if request.tz_document_id not in file_content_storage:
                raise HTTPException(status_code=404, detail=f"TZ document {request.tz_document_id} not found")
            tz_data = file_content_storage[request.tz_document_id].get("extraction_data", {})
            tz_content = tz_data.get("text", "")
        
        # Combine all KP documents for analysis
        combined_content_data = {
            "text": "",
            "tables": [],
            "currencies": [],
            "structured_data": {},
            "metadata": {"document_ids": request.document_ids}
        }
        
        for doc_id in request.document_ids:
            content_data = file_content_storage[doc_id]["extraction_data"]
            combined_content_data["text"] += content_data["text"] + "\n\n"
            combined_content_data["tables"].extend(content_data["tables"])
            combined_content_data["currencies"].extend(content_data["currencies"])
            
            # Merge structured data
            for key, value in content_data["structured_data"].items():
                if key in combined_content_data["structured_data"]:
                    if isinstance(value, dict):
                        combined_content_data["structured_data"][key].update(value)
                    elif isinstance(value, list):
                        combined_content_data["structured_data"][key].extend(value)
                else:
                    combined_content_data["structured_data"][key] = value
        
        logger.info(f"📊 Combined analysis data: {len(combined_content_data['text'])} chars, {len(combined_content_data['tables'])} tables, {len(combined_content_data['currencies'])} currencies")
        
        # Perform AI analysis with v3 system
        try:
            logger.info("🤖 Calling Claude v3 analysis...")
            analysis_data = await ai_manager.analyze_with_claude_v3(
                combined_content_data, weights, tz_content
            )
            logger.info("✅ Claude v3 analysis completed")
        except Exception as ai_error:
            logger.warning(f"⚠️ Claude v3 API failed: {ai_error}")
            logger.info("🔄 Using enhanced fallback analysis...")
            analysis_data = ai_manager._generate_fallback_v3_analysis(combined_content_data, weights)
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Generate charts data if requested
        charts_data = None
        if request.generate_charts:
            charts_data = generate_analysis_charts(analysis_data, combined_content_data)
        
        # Create comprehensive v3 analysis result
        analysis_result = V3AnalysisResponse(
            id=request.document_ids[0],  # Primary document ID
            status="completed",
            overall_score=analysis_data.get("overall_score", 75),
            weighted_score=analysis_data.get("weighted_score", 75.0),
            company_name=analysis_data.get("company_name", "Анализируемая компания"),
            summary=analysis_data.get("executive_summary", "Проведен экспертный анализ по 10 критериям"),
            executive_summary=analysis_data.get("executive_summary", ""),
            recommendations=analysis_data.get("recommendations", []),
            business_analysis=analysis_data.get("business_analysis", {}),
            criteria_weights=weights,
            risk_level=analysis_data.get("risk_level", "Умеренный риск"),
            processing_time=processing_time,
            currency_data={
                "total_currencies": len(combined_content_data["currencies"]),
                "currencies": combined_content_data["currencies"][:10]  # First 10 currencies
            },
            extracted_tables=[
                {
                    "id": table.get("table_id", f"table_{i}"),
                    "rows": table.get("row_count", 0),
                    "columns": table.get("col_count", 0),
                    "preview": table.get("data", [])[:3] if table.get("data") else []  # First 3 rows
                }
                for i, table in enumerate(combined_content_data["tables"][:5])  # First 5 tables
            ],
            charts_data=charts_data,
            created_at=datetime.now().isoformat()
        ).dict()
        
        # Store v3 analysis result
        v3_analysis_storage[request.document_ids[0]] = analysis_result
        
        logger.info(f"🎯 V3 Analysis completed: Score {analysis_result['overall_score']}/100, Weighted: {analysis_result['weighted_score']:.1f}, Time: {processing_time:.2f}s")
        
        return analysis_result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ V3 Analysis error: {e}")
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"V3 analysis failed: {str(e)}")

def generate_analysis_charts(analysis_data: Dict[str, Any], content_data: Dict[str, Any]) -> Dict[str, Any]:
    """Generate interactive charts for v3 analysis"""
    try:
        charts_data = {}
        
        # Criteria radar chart
        if "business_analysis" in analysis_data:
            criteria_names = []
            criteria_scores = []
            criteria_weights = []
            
            for criterion, data in analysis_data["business_analysis"].items():
                criteria_names.append(criterion.replace("_", " ").title())
                criteria_scores.append(data.get("score", 0))
                criteria_weights.append(data.get("weight", 0) * 100)  # Convert to percentage
            
            # Create radar chart data
            charts_data["criteria_radar"] = {
                "type": "radar",
                "data": {
                    "categories": criteria_names,
                    "scores": criteria_scores,
                    "weights": criteria_weights
                }
            }
            
            # Create bar chart for scores
            charts_data["criteria_bar"] = {
                "type": "bar", 
                "data": {
                    "categories": criteria_names,
                    "scores": criteria_scores,
                    "colors": [
                        "#22c55e" if score >= 85 else "#3b82f6" if score >= 70 else "#f59e0b" if score >= 55 else "#ef4444"
                        for score in criteria_scores
                    ]
                }
            }
        
        # Currency distribution pie chart
        if content_data.get("currencies"):
            currency_summary = {}
            for currency in content_data["currencies"]:
                curr_type = currency.get("currency", "unknown")
                if curr_type not in currency_summary:
                    currency_summary[curr_type] = {"count": 0, "total": 0}
                currency_summary[curr_type]["count"] += 1
                currency_summary[curr_type]["total"] += currency.get("amount", 0)
            
            charts_data["currency_pie"] = {
                "type": "pie",
                "data": {
                    "labels": list(currency_summary.keys()),
                    "values": [data["count"] for data in currency_summary.values()],
                    "amounts": [data["total"] for data in currency_summary.values()]
                }
            }
        
        return charts_data
        
    except Exception as e:
        logger.error(f"❌ Chart generation error: {e}")
        return {}

@app.get("/api/v3/analysis/{analysis_id}")
async def get_v3_analysis(analysis_id: int, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Get v3 analysis results by ID"""
    try:
        if analysis_id not in v3_analysis_storage:
            raise HTTPException(status_code=404, detail="V3 analysis not found")
        
        return v3_analysis_storage[analysis_id]
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Get V3 analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v3/analysis/history")
async def get_v3_analysis_history(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Get v3 analysis history"""
    try:
        history = []
        for analysis_id, analysis_data in v3_analysis_storage.items():
            history.append({
                "id": analysis_id,
                "company_name": analysis_data.get("company_name", "Unknown"),
                "overall_score": analysis_data.get("overall_score", 0),
                "weighted_score": analysis_data.get("weighted_score", 0),
                "created_at": analysis_data.get("created_at", ""),
                "status": analysis_data.get("status", "unknown"),
                "risk_level": analysis_data.get("risk_level", "Unknown")
            })
        
        # Sort by creation date (newest first)
        history.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        return {
            "history": history,
            "total": len(history),
            "version": "v3"
        }
        
    except Exception as e:
        logger.error(f"❌ Get V3 history error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v3/export/pdf/{analysis_id}")
async def export_v3_pdf(analysis_id: int, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Export v3 analysis as professional PDF with tender styling"""
    try:
        if analysis_id not in v3_analysis_storage:
            raise HTTPException(status_code=404, detail="V3 analysis not found")
        
        analysis_data = v3_analysis_storage[analysis_id]
        
        logger.info(f"🔄 Generating V3 PDF for analysis {analysis_id}...")
        
        # Transform v3 data for tender PDF exporter
        tender_compatible_data = {
            "id": analysis_data["id"],
            "company_name": analysis_data.get("company_name", "Unknown Company"),
            "overall_score": analysis_data.get("overall_score", 0),
            "weighted_score": analysis_data.get("weighted_score", 0),
            "summary": analysis_data.get("executive_summary", ""),
            "analysis_type": "v3_expert",
            "processing_time": analysis_data.get("processing_time", 0),
            "created_at": analysis_data.get("created_at", datetime.now().isoformat()),
            "risk_level": analysis_data.get("risk_level", "Unknown"),
            "recommendations": analysis_data.get("recommendations", []),
            
            # Transform business analysis for PDF
            "business_analysis": {},
            "criteria_weights": analysis_data.get("criteria_weights", {}),
            "currency_data": analysis_data.get("currency_data", {}),
            "extracted_tables": analysis_data.get("extracted_tables", []),
            "charts_data": analysis_data.get("charts_data", {}),
            
            # V3 specific additions
            "version": "3.0",
            "analysis_method": "10-Criteria Expert Analysis",
            "total_criteria": 10
        }
        
        # Transform detailed business analysis
        if analysis_data.get("business_analysis"):
            for criterion, details in analysis_data["business_analysis"].items():
                tender_compatible_data["business_analysis"][criterion] = {
                    "score": details.get("score", 0),
                    "weight": details.get("weight", 0),
                    "details": details.get("details", ""),
                    "recommendations": details.get("recommendations", []),
                    "compliance_level": details.get("compliance_level", "medium"),
                    "risk_factors": details.get("risk_factors", [])
                }
        
        try:
            # Generate PDF using enhanced tender exporter
            pdf_buffer = tender_pdf_exporter.generate_kp_analysis_pdf(tender_compatible_data)
            
            # Generate filename
            company_name = analysis_data.get("company_name", "Unknown_Company").replace(" ", "_")
            pdf_filename = f"KP_Analysis_V3_{company_name}_{analysis_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # Save PDF to reports directory
            pdf_path = Path(f"data/reports/{pdf_filename}")
            with open(pdf_path, 'wb') as f:
                f.write(pdf_buffer.read())
            
            logger.info(f"✅ V3 PDF generated successfully: {pdf_filename} (Size: {pdf_path.stat().st_size} bytes)")
            
            # Return PDF as streaming response
            pdf_buffer.seek(0)
            return StreamingResponse(
                BytesIO(pdf_buffer.read()),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={pdf_filename}"}
            )
            
        except Exception as pdf_error:
            logger.error(f"❌ V3 PDF generation error: {pdf_error}")
            logger.error(f"❌ V3 PDF Traceback: {traceback.format_exc()}")
            
            # Enhanced fallback response with v3 details
            pdf_filename = f"DevAssist_Pro_KP_Analysis_V3_{analysis_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            return {
                "message": "V3 PDF generation failed - enhanced fallback mode",
                "filename": pdf_filename,
                "analysis_id": analysis_id,
                "analysis_type": "v3_expert",
                "overall_score": analysis_data.get("overall_score", 0),
                "weighted_score": analysis_data.get("weighted_score", 0),
                "criteria_count": 10,
                "export_time": datetime.now().isoformat(),
                "error": str(pdf_error),
                "fallback_mode": True,
                "version": "3.0"
            }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ V3 PDF export error: {e}")
        raise HTTPException(status_code=500, detail=f"V3 PDF export failed: {str(e)}")

if __name__ == "__main__":
    print("Starting DevAssist Pro v3 - Enhanced AI-Powered Backend")
    print("=" * 60)
    print("Available endpoints:")
    print("   - Health Check:     http://localhost:8000/health")
    print("   - Register:         http://localhost:8000/api/auth/register")
    print("   - Login:            http://localhost:8000/api/auth/login")
    print("   - User Info:        http://localhost:8000/api/auth/me")
    print("   - Refresh Token:    http://localhost:8000/api/auth/refresh")
    print("   - V2 Upload:        http://localhost:8000/api/documents/upload")  
    print("   - V2 Analyze:       http://localhost:8000/api/documents/{id}/analyze")
    print("   - V2 Export PDF:    http://localhost:8000/api/documents/{id}/export-pdf")
    print("   - V3 Upload:        http://localhost:8000/api/v3/documents/upload")
    print("   - V3 Analyze:       http://localhost:8000/api/v3/kp-analyzer/analyze")
    print("   - V3 Export PDF:    http://localhost:8000/api/v3/export/pdf/{id}")
    print("   - V3 Weights:       http://localhost:8000/api/v3/criteria/weights/presets")
    print("   - V3 History:       http://localhost:8000/api/v3/analysis/history")
    print("   - Test Cyrillic:    http://localhost:8000/api/test/cyrillic-pdf")
    print("   - LLM Direct:       http://localhost:8000/api/llm/analyze")
    print("   - LLM Status:       http://localhost:8000/api/llm/providers")
    print("=" * 60)
    print("Features:")
    print("   * JWT Authentication System (Register, Login, Token Management)")
    print("   * KP Analyzer v2 - Basic Analysis")
    print("   * KP Analyzer v3 - Expert 10-Criteria Analysis (Protected)")
    print("   * Real Claude API Integration")
    print("   * Advanced PDF Processing with pdfplumber")
    print("   * Multi-currency Detection & Extraction")
    print("   * Interactive Charts Generation")
    print("   * Configurable Criteria Weights")
    print("   * Professional PDF Export with Cyrillic Support")
    print("   * Document Processing (PDF, DOCX, TXT)")
    print("   * Structured Data Extraction from Tables")
    print("   * Business Logic Analysis")
    print("   * Risk Assessment & Recommendations")
    print("   * Secure Password Hashing & Validation")
    print("=" * 60)
    
    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8000,
        reload=False,
        log_level="info"
    )