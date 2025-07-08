"""
Text Extractor для Documents Service
Извлечение текста из различных форматов файлов (портировано из старого приложения)
"""
import os
import asyncio
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import PyPDF2
import docx
from io import BytesIO
import aiofiles

logger = logging.getLogger(__name__)

class TextExtractor:
    """Класс для извлечения текста из документов"""
    
    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".txt"]
    
    async def extract_text_async(self, file_path: Path) -> str:
        """
        Асинхронное извлечение текста из файла
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            str: Извлеченный текст
        """
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == ".pdf":
                return await self._extract_from_pdf(file_path)
            elif file_extension == ".docx":
                return await self._extract_from_docx(file_path)
            elif file_extension == ".txt":
                return await self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unknown file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    def extract_text_sync(self, file_path: Path) -> str:
        """
        Синхронное извлечение текста из файла (портировано из старого приложения)
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            str: Извлеченный текст
        """
        file_extension = file_path.suffix.lower()
        
        if file_extension == ".pdf":
            return self._extract_text_from_pdf_sync(file_path)
        elif file_extension == ".docx":
            return self._extract_text_from_docx_sync(file_path)
        elif file_extension == ".txt":
            return self._extract_text_from_txt_sync(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат файла: {file_extension}")
            return ""
    
    async def _extract_from_pdf(self, file_path: Path) -> str:
        """Асинхронное извлечение текста из PDF"""
        def extract_pdf():
            return self._extract_text_from_pdf_sync(file_path)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract_pdf)
    
    async def _extract_from_docx(self, file_path: Path) -> str:
        """Асинхронное извлечение текста из DOCX"""
        def extract_docx():
            return self._extract_text_from_docx_sync(file_path)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract_docx)
    
    async def _extract_from_txt(self, file_path: Path) -> str:
        """Асинхронное извлечение текста из TXT"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                return await f.read()
        except UnicodeDecodeError:
            # Попробовать другие кодировки
            for encoding in ["cp1251", "latin-1"]:
                try:
                    async with aiofiles.open(file_path, "r", encoding=encoding) as f:
                        return await f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file with any supported encoding")
    
    def _extract_text_from_pdf_sync(self, pdf_path: Path) -> str:
        """
        Извлекает текст из PDF-файла (портировано из старого приложения)
        
        Args:
            pdf_path: Путь к PDF-файлу
            
        Returns:
            str: Извлеченный текст
        """
        try:
            text = ""
            with open(pdf_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из PDF {pdf_path}: {e}")
            return ""
    
    def _extract_text_from_docx_sync(self, docx_path: Path) -> str:
        """
        Извлекает текст из DOCX-файла (портировано из старого приложения)
        
        Args:
            docx_path: Путь к DOCX-файлу
            
        Returns:
            str: Извлеченный текст
        """
        try:
            doc = docx.Document(docx_path)
            text_parts = []
            
            # Извлечение текста из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOCX {docx_path}: {e}")
            return ""
    
    def _extract_text_from_txt_sync(self, txt_path: Path) -> str:
        """
        Извлекает текст из TXT-файла (портировано из старого приложения)
        
        Args:
            txt_path: Путь к TXT-файлу
            
        Returns:
            str: Извлеченный текст
        """
        try:
            # Попробовать UTF-8 сначала
            with open(txt_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Попробовать другие популярные кодировки
            for encoding in ["cp1251", "latin-1", "cp866"]:
                try:
                    with open(txt_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            logger.error(f"Не удалось декодировать текстовый файл {txt_path}")
            return ""
        except Exception as e:
            logger.error(f"Ошибка при чтении текстового файла {txt_path}: {e}")
            return ""
    
    def get_document_info(self, file_path: Path) -> Dict[str, Any]:
        """
        Получение информации о документе
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Dict с информацией о документе
        """
        try:
            stat = file_path.stat()
            
            info = {
                "filename": file_path.name,
                "file_size": stat.st_size,
                "file_type": file_path.suffix.lower(),
                "created_at": stat.st_ctime,
                "modified_at": stat.st_mtime,
                "is_supported": file_path.suffix.lower() in self.supported_formats
            }
            
            # Дополнительная информация для PDF
            if file_path.suffix.lower() == ".pdf":
                try:
                    with open(file_path, "rb") as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        info.update({
                            "page_count": len(pdf_reader.pages),
                            "has_metadata": bool(pdf_reader.metadata),
                            "is_encrypted": pdf_reader.is_encrypted
                        })
                        
                        if pdf_reader.metadata:
                            info["metadata"] = {
                                "title": pdf_reader.metadata.get("/Title", ""),
                                "author": pdf_reader.metadata.get("/Author", ""),
                                "subject": pdf_reader.metadata.get("/Subject", ""),
                                "creator": pdf_reader.metadata.get("/Creator", "")
                            }
                except Exception as e:
                    logger.warning(f"Could not extract PDF metadata from {file_path}: {e}")
            
            # Дополнительная информация для DOCX
            elif file_path.suffix.lower() == ".docx":
                try:
                    doc = docx.Document(file_path)
                    info.update({
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables)
                    })
                    
                    # Метаданные документа
                    if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                        info["metadata"] = {
                            "title": doc.core_properties.title,
                            "author": doc.core_properties.author,
                            "subject": doc.core_properties.subject,
                            "created": doc.core_properties.created,
                            "modified": doc.core_properties.modified
                        }
                except Exception as e:
                    logger.warning(f"Could not extract DOCX metadata from {file_path}: {e}")
            
            return info
            
        except Exception as e:
            logger.error(f"Failed to get document info for {file_path}: {e}")
            return {
                "filename": file_path.name,
                "error": str(e),
                "is_supported": False
            }
    
    def validate_file(self, file_path: Path) -> bool:
        """
        Валидация файла
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            bool: True если файл валиден
        """
        if not file_path.exists():
            return False
        
        if file_path.suffix.lower() not in self.supported_formats:
            return False
        
        # Проверка размера файла (максимум 50MB)
        max_size = 50 * 1024 * 1024  # 50MB
        if file_path.stat().st_size > max_size:
            logger.warning(f"File {file_path} is too large: {file_path.stat().st_size} bytes")
            return False
        
        # Базовая проверка содержимого
        try:
            if file_path.suffix.lower() == ".pdf":
                with open(file_path, "rb") as f:
                    PyPDF2.PdfReader(f)
            elif file_path.suffix.lower() == ".docx":
                docx.Document(file_path)
            elif file_path.suffix.lower() == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    f.read(1024)  # Проверить первый кб
            
            return True
            
        except Exception as e:
            logger.error(f"File validation failed for {file_path}: {e}")
            return False