"""
Расширенный процессор документов для v3
Поддержка advanced extraction с таблицами и валютами
"""
import logging
import asyncio
from typing import Dict, Any, List, Optional
from io import BytesIO
import json

# Document processing imports
import PyPDF2
import docx
import pdfplumber
import re

logger = logging.getLogger(__name__)

class V3DocumentProcessor:
    """Продвинутый процессор документов с поддержкой структурированных данных"""
    
    def __init__(self):
        self.currency_patterns = {
            'som': [
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:сом|som|KGS)',
                r'(?:сом|som|KGS)\s*(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
            ],
            'ruble': [
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:руб|rub|RUB|₽)',
                r'(?:руб|rub|RUB|₽)\s*(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
            ],
            'dollar': [
                r'\$(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:USD|dollar)',
            ],
            'euro': [
                r'€(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)',
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:EUR|euro)',
            ],
            'tenge': [
                r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)\s*(?:тенге|tenge|KZT)',
            ]
        }
    
    async def extract_advanced_content(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Расширенное извлечение контента с поддержкой таблиц и валют"""
        try:
            logger.info(f"🔍 Starting advanced extraction for {filename}")
            
            extraction_result = {
                "text": "",
                "tables": [],
                "currencies": [],
                "structured_data": {},
                "metadata": {
                    "filename": filename,
                    "extraction_method": "v3_advanced",
                    "processing_timestamp": str(asyncio.get_event_loop().time())
                }
            }
            
            # Определяем тип файла и обрабатываем соответственно
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                extraction_result = await self._process_pdf_advanced(file_content, extraction_result)
            elif file_ext in ['docx', 'doc']:
                extraction_result = await self._process_docx_advanced(file_content, extraction_result)
            elif file_ext == 'txt':
                extraction_result = await self._process_txt_advanced(file_content, extraction_result)
            else:
                # Fallback to basic text extraction
                extraction_result["text"] = file_content.decode('utf-8', errors='ignore')
            
            # Извлекаем валюты из текста
            extraction_result["currencies"] = self._extract_currencies(extraction_result["text"])
            
            # Структурируем финансовые данные
            extraction_result["structured_data"] = self._structure_financial_data(
                extraction_result["text"], 
                extraction_result["tables"],
                extraction_result["currencies"]
            )
            
            logger.info(f"✅ Advanced extraction complete: {len(extraction_result['text'])} chars, {len(extraction_result['tables'])} tables, {len(extraction_result['currencies'])} currencies")
            
            return extraction_result
            
        except Exception as e:
            logger.error(f"❌ Advanced extraction error: {e}")
            # Fallback to basic extraction
            return {
                "text": file_content.decode('utf-8', errors='ignore'),
                "tables": [],
                "currencies": [],
                "structured_data": {},
                "metadata": {
                    "error": str(e),
                    "extraction_method": "fallback_basic"
                }
            }
    
    async def _process_pdf_advanced(self, file_content: bytes, result: Dict) -> Dict:
        """Продвинутая обработка PDF с извлечением таблиц"""
        try:
            # Попытка использовать pdfplumber для лучшего извлечения
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                full_text = ""
                tables = []
                
                for page_num, page in enumerate(pdf.pages):
                    # Извлечение текста
                    page_text = page.extract_text() or ""
                    full_text += page_text + "\n"
                    
                    # Извлечение таблиц
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables or []):
                        if table and len(table) > 1:  # Валидная таблица с заголовками
                            tables.append({
                                "page": page_num + 1,
                                "table_id": f"pdf_table_p{page_num}_{table_idx}",
                                "data": table,
                                "row_count": len(table),
                                "col_count": len(table[0]) if table else 0,
                                "extraction_method": "pdfplumber"
                            })
                
                result["text"] = full_text
                result["tables"] = tables
                
        except Exception as pdfplumber_error:
            logger.warning(f"⚠️ PDFPlumber failed, falling back to PyPDF2: {pdfplumber_error}")
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                result["text"] = text
                result["metadata"]["fallback_used"] = "PyPDF2"
            except Exception as pypdf2_error:
                logger.error(f"❌ PyPDF2 also failed: {pypdf2_error}")
                result["text"] = "PDF extraction failed"
                result["metadata"]["error"] = str(pypdf2_error)
        
        return result
    
    async def _process_docx_advanced(self, file_content: bytes, result: Dict) -> Dict:
        """Продвинутая обработка DOCX с извлечением таблиц"""
        try:
            doc = docx.Document(BytesIO(file_content))
            
            # Извлечение текста из параграфов
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            result["text"] = text
            
            # Извлечение таблиц
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        "table_id": f"docx_table_{table_idx}",
                        "data": table_data,
                        "row_count": len(table_data),
                        "col_count": len(table_data[0]) if table_data else 0,
                        "extraction_method": "python-docx"
                    })
            
            result["tables"] = tables
            
        except Exception as e:
            logger.error(f"❌ DOCX processing error: {e}")
            result["text"] = "DOCX extraction failed"
            result["metadata"]["error"] = str(e)
        
        return result
    
    async def _process_txt_advanced(self, file_content: bytes, result: Dict) -> Dict:
        """Обработка текстовых файлов с определением кодировки"""
        try:
            # Попытка определить кодировку
            import chardet
            detected = chardet.detect(file_content)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)
            
            logger.info(f"Detected encoding: {encoding} (confidence: {confidence:.2f})")
            
            # Декодируем с определенной кодировкой
            text = file_content.decode(encoding, errors='ignore')
            result["text"] = text
            result["metadata"]["encoding"] = encoding
            result["metadata"]["encoding_confidence"] = confidence
            
        except Exception as e:
            logger.error(f"❌ TXT processing error: {e}")
            # Fallback to UTF-8
            result["text"] = file_content.decode('utf-8', errors='ignore')
            result["metadata"]["error"] = str(e)
        
        return result
    
    def _extract_currencies(self, text: str) -> List[Dict[str, Any]]:
        """Извлечение валют с суммами из текста"""
        currencies = []
        
        for currency_type, pattern_list in self.currency_patterns.items():
            for pattern in pattern_list:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.group(1)
                    # Очистка и конвертация суммы
                    amount_clean = re.sub(r'[^\d.,]', '', amount_str)
                    amount_clean = amount_clean.replace(',', '.')
                    
                    try:
                        amount = float(amount_clean)
                        currencies.append({
                            "currency": currency_type,
                            "amount": amount,
                            "formatted": f"{amount:,.2f}",
                            "original_text": match.group(0),
                            "position": match.start(),
                            "context": text[max(0, match.start()-50):match.end()+50]
                        })
                    except ValueError:
                        # Не удалось распарсить число
                        continue
        
        # Сортировка по позиции в тексте
        currencies.sort(key=lambda x: x["position"])
        
        return currencies
    
    def _structure_financial_data(self, text: str, tables: List[Dict], currencies: List[Dict]) -> Dict[str, Any]:
        """Структурирование финансовых данных из текста и таблиц"""
        structured_data = {
            "budget_breakdown": {},
            "timeline_data": {},
            "team_structure": {},
            "technical_requirements": {},
            "currency_summary": {}
        }
        
        # Анализ таблиц на предмет финансовых данных
        for table in tables:
            table_data = table.get("data", [])
            if not table_data or len(table_data) < 2:
                continue
            
            headers = [str(cell).lower() if cell else "" for cell in table_data[0]]
            
            # Поиск финансовых таблиц
            if any(keyword in " ".join(headers) for keyword in ["стоимость", "цена", "бюджет", "сумма", "cost", "price"]):
                budget_data = {}
                for row in table_data[1:]:
                    if len(row) >= 2:
                        item = str(row[0]).strip()
                        amount_str = str(row[1]).strip()
                        
                        # Попытка извлечь числовое значение
                        amount_match = re.search(r'(\d+(?:\s?\d{3})*(?:[,.]?\d{2})?)', amount_str)
                        if amount_match:
                            try:
                                numeric_amount = float(amount_match.group(1).replace(' ', '').replace(',', '.'))
                                budget_data[item] = {
                                    "amount": numeric_amount,
                                    "formatted": f"{numeric_amount:,.2f}",
                                    "original": amount_str
                                }
                            except ValueError:
                                continue
                
                if budget_data:
                    structured_data["budget_breakdown"].update(budget_data)
        
        # Суммирование валют по типам
        currency_summary = {}
        for curr in currencies:
            curr_type = curr.get("currency", "unknown")
            if curr_type not in currency_summary:
                currency_summary[curr_type] = {
                    "count": 0,
                    "total_amount": 0,
                    "amounts": []
                }
            
            currency_summary[curr_type]["count"] += 1
            currency_summary[curr_type]["total_amount"] += curr.get("amount", 0)
            currency_summary[curr_type]["amounts"].append(curr.get("amount", 0))
        
        structured_data["currency_summary"] = currency_summary
        
        # Анализ временных данных из текста
        timeline_keywords = ["этап", "месяц", "неделя", "день", "срок", "deadline", "schedule"]
        timeline_matches = []
        for keyword in timeline_keywords:
            pattern = rf'{keyword}[:\s]*(\d+(?:\s?\d+)*)'
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                timeline_matches.append({
                    "keyword": keyword,
                    "value": match.group(1),
                    "context": text[max(0, match.start()-30):match.end()+30]
                })
        
        structured_data["timeline_data"] = {
            "matches": timeline_matches[:10],  # Первые 10 совпадений
            "total_matches": len(timeline_matches)
        }
        
        return structured_data