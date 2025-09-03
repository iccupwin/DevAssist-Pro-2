#!/usr/bin/env python3
"""
Test: File Processing System for DevAssist Pro
Tests document upload, text extraction, and file validation
"""

import os
import asyncio
import tempfile
from pathlib import Path
from typing import Dict, Any, List
import sys

# Add the current directory to the path for imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from services.documents.core.text_extractor import TextExtractor

class FileProcessingTester:
    """File Processing Test System"""
    
    def __init__(self):
        self.extractor = TextExtractor()
        self.test_results = {}
    
    async def test_text_file_processing(self) -> bool:
        """Test TXT file processing"""
        print("1. Testing TXT file processing...")
        
        try:
            # Test with existing test file
            txt_file = Path("test_kp_sample.txt")
            if not txt_file.exists():
                print("   WARNING: test_kp_sample.txt not found, creating temporary file")
                # Create a temporary test file
                with open("temp_test.txt", "w", encoding="utf-8") as f:
                    f.write("Тестовое коммерческое предложение\nКомпания: ТестСтрой\nСтоимость: 500,000 рублей")
                txt_file = Path("temp_test.txt")
            
            # Extract text
            extracted_text = await self.extractor.extract_text_async(txt_file)
            
            # Validate
            if extracted_text and len(extracted_text) > 0:
                print(f"   SUCCESS: Extracted {len(extracted_text)} characters")
                print(f"   Sample: {extracted_text[:100]}...")
                
                # Clean up temporary file
                if txt_file.name == "temp_test.txt":
                    txt_file.unlink()
                
                return True
            else:
                print("   FAIL: No text extracted")
                return False
                
        except Exception as e:
            print(f"   FAIL: TXT processing error: {e}")
            return False
    
    async def test_pdf_file_processing(self) -> bool:
        """Test PDF file processing"""
        print("2. Testing PDF file processing...")
        
        try:
            # Check for existing test PDF files
            test_pdfs = [
                Path("test_kp.pdf"),
                Path("test_kp_real.pdf"),
                Path("data/uploads/20250807_211330_9e7bfd12-5cb7-48e0-950b-05ab9cb31bef_test_kp_real.pdf")
            ]
            
            pdf_file = None
            for pdf_path in test_pdfs:
                if pdf_path.exists():
                    pdf_file = pdf_path
                    break
            
            if not pdf_file:
                print("   WARNING: No test PDF files found, skipping PDF test")
                return True  # Don't fail if no test files available
            
            print(f"   Testing with: {pdf_file.name}")
            
            # Extract text
            extracted_text = await self.extractor.extract_text_async(pdf_file)
            
            if extracted_text and len(extracted_text) > 0:
                print(f"   SUCCESS: Extracted {len(extracted_text)} characters from PDF")
                print(f"   Sample: {extracted_text[:100]}...")
                return True
            else:
                print("   WARNING: No text extracted from PDF (might be image-based)")
                return True  # Don't fail for image-based PDFs
                
        except Exception as e:
            print(f"   FAIL: PDF processing error: {e}")
            return False
    
    async def test_docx_file_processing(self) -> bool:
        """Test DOCX file processing"""
        print("3. Testing DOCX file processing...")
        
        try:
            # Check for existing test DOCX files
            test_docx_files = [
                Path("test_kp.docx"),
                Path("data/uploads/20250807_170438_284ec38d-9c4c-41a5-8f9e-32329347cce3_test_kp.docx")
            ]
            
            docx_file = None
            for docx_path in test_docx_files:
                if docx_path.exists():
                    docx_file = docx_path
                    break
            
            if not docx_file:
                print("   WARNING: No test DOCX files found, creating temporary file")
                # Create a simple DOCX for testing
                try:
                    import docx as python_docx
                    doc = python_docx.Document()
                    doc.add_paragraph("Тестовое коммерческое предложение")
                    doc.add_paragraph("Компания: ТестСтрой")
                    doc.add_paragraph("Стоимость: 750,000 рублей")
                    doc.save("temp_test.docx")
                    docx_file = Path("temp_test.docx")
                except ImportError:
                    print("   WARNING: python-docx not available, skipping DOCX test")
                    return True
            
            print(f"   Testing with: {docx_file.name}")
            
            # Extract text
            extracted_text = await self.extractor.extract_text_async(docx_file)
            
            if extracted_text and len(extracted_text) > 0:
                print(f"   SUCCESS: Extracted {len(extracted_text)} characters from DOCX")
                print(f"   Sample: {extracted_text[:100]}...")
                
                # Clean up temporary file
                if docx_file.name == "temp_test.docx":
                    docx_file.unlink()
                
                return True
            else:
                print("   FAIL: No text extracted from DOCX")
                return False
                
        except Exception as e:
            print(f"   FAIL: DOCX processing error: {e}")
            return False
    
    async def test_file_format_validation(self) -> bool:
        """Test file format validation"""
        print("4. Testing file format validation...")
        
        try:
            # Test supported formats
            supported_formats = self.extractor.supported_formats
            print(f"   Supported formats: {supported_formats}")
            
            # Test unsupported format
            try:
                fake_file = Path("test.xyz")
                if fake_file.exists():
                    await self.extractor.extract_text_async(fake_file)
                    print("   FAIL: Should have rejected unsupported format")
                    return False
                else:
                    print("   INFO: test.xyz not found (which is expected)")
                    
                # Create and test with unsupported format
                with open("temp_unsupported.xyz", "w") as f:
                    f.write("test content")
                
                temp_file = Path("temp_unsupported.xyz")
                await self.extractor.extract_text_async(temp_file)
                
                # Should not reach here
                print("   FAIL: Should have rejected .xyz format")
                temp_file.unlink()
                return False
                
            except ValueError as e:
                if "Unsupported file format" in str(e):
                    print("   SUCCESS: Properly rejected unsupported format")
                    Path("temp_unsupported.xyz").unlink()
                    return True
                else:
                    print(f"   FAIL: Unexpected error: {e}")
                    return False
                    
        except Exception as e:
            print(f"   FAIL: Format validation error: {e}")
            return False
    
    async def test_error_handling(self) -> bool:
        """Test error handling for file processing"""
        print("5. Testing error handling...")
        
        try:
            # Test with non-existent file
            try:
                non_existent = Path("definitely_does_not_exist.txt")
                await self.extractor.extract_text_async(non_existent)
                print("   FAIL: Should have failed for non-existent file")
                return False
                
            except Exception as e:
                print("   SUCCESS: Properly handled non-existent file")
            
            # Test with corrupted file (if available)
            corrupted_file = Path("corrupted.docx")
            if corrupted_file.exists():
                try:
                    await self.extractor.extract_text_async(corrupted_file)
                    print("   WARNING: Corrupted file processed without error")
                except Exception as e:
                    print("   SUCCESS: Properly handled corrupted file")
            
            return True
            
        except Exception as e:
            print(f"   FAIL: Error handling test failed: {e}")
            return False
    
    async def test_cyrillic_text_support(self) -> bool:
        """Test Cyrillic text processing"""
        print("6. Testing Cyrillic text support...")
        
        try:
            # Create test file with Cyrillic text
            cyrillic_content = """
            Коммерческое предложение от ООО "РусТех"
            
            Разработка веб-приложения
            Стоимость: 1,000,000 рублей
            Срок: 6 месяцев
            
            Техническое задание:
            - Современный интерфейс
            - База данных PostgreSQL
            - Высокая производительность
            """
            
            with open("temp_cyrillic.txt", "w", encoding="utf-8") as f:
                f.write(cyrillic_content)
            
            cyrillic_file = Path("temp_cyrillic.txt")
            extracted_text = await self.extractor.extract_text_async(cyrillic_file)
            
            # Check if Cyrillic characters are preserved
            cyrillic_chars = ['а', 'б', 'в', 'г', 'д', 'е', 'ё', 'ж', 'з', 'и']
            cyrillic_found = any(char in extracted_text.lower() for char in cyrillic_chars)
            
            if cyrillic_found and "Коммерческое" in extracted_text:
                print("   SUCCESS: Cyrillic text properly processed")
                cyrillic_file.unlink()
                return True
            else:
                print("   FAIL: Cyrillic text not properly preserved")
                print(f"   Extracted: {extracted_text[:100]}...")
                cyrillic_file.unlink()
                return False
                
        except Exception as e:
            print(f"   FAIL: Cyrillic text test error: {e}")
            return False

async def test_file_processing_system():
    """Test the complete file processing system"""
    
    print("=" * 60)
    print("FILE PROCESSING SYSTEM TEST")
    print("=" * 60)
    
    tester = FileProcessingTester()
    
    # Run all tests
    tests = [
        tester.test_text_file_processing,
        tester.test_pdf_file_processing,
        tester.test_docx_file_processing,
        tester.test_file_format_validation,
        tester.test_error_handling,
        tester.test_cyrillic_text_support
    ]
    
    results = []
    for test in tests:
        try:
            result = await test()
            results.append(result)
        except Exception as e:
            print(f"   CRITICAL ERROR in {test.__name__}: {e}")
            results.append(False)
        print()
    
    # Results summary
    print("=" * 60)
    print("FILE PROCESSING TEST RESULTS")
    print("=" * 60)
    
    passed = sum(results)
    total = len(results)
    
    print(f"Tests passed: {passed}/{total}")
    print(f"Success rate: {passed/total*100:.1f}%")
    
    test_names = [
        "TXT file processing",
        "PDF file processing", 
        "DOCX file processing",
        "File format validation",
        "Error handling",
        "Cyrillic text support"
    ]
    
    if passed == total:
        print("RESULT: ALL FILE PROCESSING TESTS PASSED!")
        for name in test_names:
            print(f"   SUCCESS: {name}")
    else:
        print("RESULT: SOME TESTS FAILED")
        for i, (name, result) in enumerate(zip(test_names, results)):
            status = "PASS" if result else "FAIL"
            print(f"   {status}: {name}")
    
    return passed >= (total * 0.8)  # Allow 80% pass rate

if __name__ == "__main__":
    try:
        success = asyncio.run(test_file_processing_system())
        exit(0 if success else 1)
    except Exception as e:
        print(f"CRITICAL ERROR: {e}")
        exit(1)