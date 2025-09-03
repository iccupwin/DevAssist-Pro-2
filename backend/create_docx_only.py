#!/usr/bin/env python3
"""Создание тестового DOCX файла"""

from docx import Document

def create_test_docx():
    """Создает тестовый DOCX файл"""
    try:
        filename = "test_kp.docx"
        doc = Document()
        
        # Добавляем заголовок
        doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', 0)
        
        # Добавляем параграфы
        doc.add_paragraph('Компания: ООО "ТехноСтрой"')
        doc.add_paragraph('Дата: 04.08.2025')
        doc.add_paragraph('')
        
        doc.add_heading('1. ОПИСАНИЕ ПРОЕКТА', level=1)
        doc.add_paragraph('Разработка информационной системы для управления строительными проектами.')
        
        doc.add_heading('2. ТЕХНОЛОГИИ', level=1)
        doc.add_paragraph('- Backend: Python, FastAPI')
        doc.add_paragraph('- Frontend: React, TypeScript')
        doc.add_paragraph('- База данных: PostgreSQL')
        doc.add_paragraph('- Инфраструктура: Docker, Kubernetes')
        
        doc.add_heading('3. КОМАНДА', level=1)
        doc.add_paragraph('- Руководитель проекта: 1 человек')
        doc.add_paragraph('- Backend разработчики: 3 человека')
        doc.add_paragraph('- Frontend разработчики: 2 человека')
        doc.add_paragraph('- DevOps инженер: 1 человек')
        
        doc.add_heading('4. СТОИМОСТЬ', level=1)
        doc.add_paragraph('Общая стоимость проекта: 5 000 000 рублей')
        doc.add_paragraph('- Проектирование: 500 000 руб')
        doc.add_paragraph('- Разработка: 3 500 000 руб')
        doc.add_paragraph('- Тестирование: 500 000 руб')
        doc.add_paragraph('- Внедрение и поддержка: 500 000 руб')
        
        doc.save(filename)
        print(f"DOCX файл создан: {filename}")
        return filename
    except Exception as e:
        print(f"Ошибка создания DOCX: {e}")
        return None

if __name__ == "__main__":
    create_test_docx()
    print("DOCX файл готов!")