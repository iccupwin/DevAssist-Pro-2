#!/usr/bin/env python3
"""Создание тестовых PDF и DOCX файлов для проверки извлечения текста"""

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
import os

def create_test_pdf():
    """Создает тестовый PDF файл"""
    try:
        filename = "test_kp.pdf"
        c = canvas.Canvas(filename, pagesize=letter)
        
        # Добавляем текст
        c.drawString(100, 750, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
        c.drawString(100, 720, "")
        c.drawString(100, 700, "Компания: ООО 'ТехноСтрой'")
        c.drawString(100, 680, "Дата: 04.08.2025")
        c.drawString(100, 650, "")
        c.drawString(100, 630, "1. ОПИСАНИЕ ПРОЕКТА")
        c.drawString(100, 610, "Разработка информационной системы управления")
        c.drawString(100, 590, "строительными проектами.")
        c.drawString(100, 560, "")
        c.drawString(100, 540, "2. ТЕХНОЛОГИИ")
        c.drawString(100, 520, "- Backend: Python, FastAPI")
        c.drawString(100, 500, "- Frontend: React, TypeScript")
        c.drawString(100, 480, "- База данных: PostgreSQL")
        c.drawString(100, 450, "")
        c.drawString(100, 430, "3. СТОИМОСТЬ")
        c.drawString(100, 410, "Общая стоимость: 5 000 000 рублей")
        c.drawString(100, 390, "- Разработка: 3 500 000 руб")
        c.drawString(100, 370, "- Тестирование: 500 000 руб")
        c.drawString(100, 350, "- Внедрение: 500 000 руб")
        
        c.save()
        print(f"PDF файл создан: {filename}")
        return filename
    except Exception as e:
        print(f"Ошибка создания PDF: {e}")
        return None

def create_test_docx():
    """Создает тестовый DOCX файл"""
    try:
        filename = "test_kp.docx"
        doc = Document()
        
        # Добавляем заголовок
        doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', 0)
        
        # Добавляем параграфы
        doc.add_paragraph('Компания: ООО "ТехноСтрой"')
        doc.add_paragraph('Дата: 04.08.2025')
        doc.add_paragraph('')
        
        doc.add_heading('1. ОПИСАНИЕ ПРОЕКТА', level=1)
        doc.add_paragraph('Разработка информационной системы для управления строительными проектами.')
        
        doc.add_heading('2. ТЕХНОЛОГИИ', level=1)
        doc.add_paragraph('- Backend: Python, FastAPI')
        doc.add_paragraph('- Frontend: React, TypeScript')
        doc.add_paragraph('- База данных: PostgreSQL')
        doc.add_paragraph('- Инфраструктура: Docker, Kubernetes')
        
        doc.add_heading('3. КОМАНДА', level=1)
        doc.add_paragraph('- Руководитель проекта: 1 человек')
        doc.add_paragraph('- Backend разработчики: 3 человека')
        doc.add_paragraph('- Frontend разработчики: 2 человека')
        doc.add_paragraph('- DevOps инженер: 1 человек')
        
        doc.add_heading('4. СТОИМОСТЬ', level=1)
        doc.add_paragraph('Общая стоимость проекта: 5 000 000 рублей')
        doc.add_paragraph('- Проектирование: 500 000 руб')
        doc.add_paragraph('- Разработка: 3 500 000 руб')
        doc.add_paragraph('- Тестирование: 500 000 руб')
        doc.add_paragraph('- Внедрение и поддержка: 500 000 руб')
        
        doc.add_heading('5. ГАРАНТИИ', level=1)
        doc.add_paragraph('- Гарантийная поддержка: 12 месяцев')
        doc.add_paragraph('- Исправление критических ошибок: в течение 24 часов')
        doc.add_paragraph('- Обучение пользователей включено в стоимость')
        
        doc.save(filename)
        print(f"DOCX файл создан: {filename}")
        return filename
    except Exception as e:
        print(f"Ошибка создания DOCX: {e}")
        return None

if __name__ == "__main__":
    create_test_docx()  # PDF создание может потребовать дополнительные библиотеки
    print("Готово!")